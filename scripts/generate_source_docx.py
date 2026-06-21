"""Generate a DOCX with all the source code for the Planned app."""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("/home/z/my-project/public/Planned_Source_Code.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# Base style
style = doc.styles["Normal"]
style.font.name = "Consolas"
style.font.size = Pt(8)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

# Files to include in order
FILES = [
    ("package.json", "/home/z/my-project/package.json"),
    (".env", "/home/z/my-project/.env"),
    ("prisma/schema.prisma", "/home/z/my-project/prisma/schema.prisma"),
    ("scripts/seed.ts", "/home/z/my-project/scripts/seed.ts"),
    ("src/lib/types.ts", "/home/z/my-project/src/lib/types.ts"),
    ("src/lib/phrases.ts", "/home/z/my-project/src/lib/phrases.ts"),
    ("src/lib/db.ts", "/home/z/my-project/src/lib/db.ts"),
    ("src/lib/store.ts", "/home/z/my-project/src/lib/store.ts"),
    ("src/lib/db-queries.ts", "/home/z/my-project/src/lib/db-queries.ts"),
    ("src/lib/store-hydration.ts", "/home/z/my-project/src/lib/store-hydration.ts"),
    ("src/lib/mutations.ts", "/home/z/my-project/src/lib/mutations.ts"),
    ("src/app/layout.tsx", "/home/z/my-project/src/app/layout.tsx"),
    ("src/app/api/state/route.ts", "/home/z/my-project/src/app/api/state/route.ts"),
    ("src/app/api/mutations/route.ts", "/home/z/my-project/src/app/api/mutations/route.ts"),
    ("src/app/globals.css", "/home/z/my-project/src/app/globals.css"),
    ("src/app/page.tsx", "/home/z/my-project/src/app/page.tsx"),
    ("src/components/avatar.tsx", "/home/z/my-project/src/components/avatar.tsx"),
    ("src/components/modals.tsx", "/home/z/my-project/src/components/modals.tsx"),
    ("src/components/child-dashboard.tsx", "/home/z/my-project/src/components/child-dashboard.tsx"),
    ("src/components/goals.tsx", "/home/z/my-project/src/components/goals.tsx"),
    ("src/components/charts.tsx", "/home/z/my-project/src/components/charts.tsx"),
    ("src/components/parent-actions.tsx", "/home/z/my-project/src/components/parent-actions.tsx"),
    ("src/components/theme-switcher.tsx", "/home/z/my-project/src/components/theme-switcher.tsx"),
    ("src/components/parent-quote-editor.tsx", "/home/z/my-project/src/components/parent-quote-editor.tsx"),
    ("src/components/family-theme-footer.tsx", "/home/z/my-project/src/components/family-theme-footer.tsx"),
]

# ---- COVER ----
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("PLANNED")
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(201, 168, 76)
r.bold = True
cover.paragraph_format.space_before = Pt(80)
cover.paragraph_format.space_after = Pt(0)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Family Savings & Allowance App — Complete Source Code")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(120, 120, 120)
r.italic = True
sub.paragraph_format.space_after = Pt(40)

# Quick start
qs_heading = doc.add_paragraph()
r = qs_heading.add_run("Quick Start")
r.font.size = Pt(16)
r.bold = True
r.font.color.rgb = RGBColor(40, 40, 40)
r.font.name = "Calibri"

qs = doc.add_paragraph()
qs.paragraph_format.line_spacing = 1.5
r = qs.add_run(
    "1. Create a new Next.js project\n"
    "2. Copy these files into the project (preserving the folder structure)\n"
    "3. Run: bun install  (or npm install)\n"
    "4. Set up the database: bun run db:push\n"
    "5. Seed demo data: bun run scripts/seed.ts\n"
    "6. Start: bun run dev\n"
    "7. Open http://localhost:3000"
)
r.font.size = Pt(11)
r.font.name = "Calibri"

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(20)
r = note.add_run(
    "Note: This document contains 25 source files (~8,800 lines). "
    "The shadcn/ui component library (src/components/ui/) is NOT included — "
    "those 40+ files are boilerplate. Generate them via: npx shadcn-ui@latest init"
)
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(140, 140, 140)
r.italic = True
r.font.name = "Calibri"

doc.add_page_break()

# ---- FILE INDEX ----
idx_heading = doc.add_paragraph()
r = idx_heading.add_run("File Index")
r.font.size = Pt(16)
r.bold = True
r.font.color.rgb = RGBColor(40, 40, 40)
r.font.name = "Calibri"

for i, (rel_path, _) in enumerate(FILES, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(f"{i}. {rel_path}")
    r.font.size = Pt(10)
    r.font.name = "Calibri"

doc.add_page_break()

# ---- EACH FILE ----
for rel_path, abs_path in FILES:
    # File heading
    heading = doc.add_paragraph()
    r = heading.add_run(f"=== {rel_path} ===")
    r.font.size = Pt(12)
r.bold = True
r.font.color.rgb = RGBColor(201, 168, 76)
r.font.name = "Consolas"
heading.paragraph_format.space_before = Pt(12)
heading.paragraph_format.space_after = Pt(6)

# Add a bottom border to the heading
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
pPr = heading._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "4")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "C9A84C")
pBdr.append(bottom)
pPr.append(pBdr)

# Read file content
try:
    with open(abs_path, "r", encoding="utf-8") as f:
        content = f.read()
except Exception as e:
    content = f"Error reading file: {e}"

# Add content as a single paragraph with the code
# Split into lines and add each as a run in one paragraph for compactness
lines = content.split("\n")
# Add code in chunks to avoid paragraph limits
chunk_size = 200
for chunk_start in range(0, len(lines), chunk_size):
    chunk = "\n".join(lines[chunk_start:chunk_start + chunk_size])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(chunk)
    r.font.size = Pt(7.5)
    r.font.name = "Consolas"

# Page break between files
doc.add_page_break()

doc.save(str(OUT))
print(f"DOCX saved: {OUT}")
print(f"Size: {OUT.stat().st_size:,} bytes")
print(f"Files included: {len(FILES)}")
