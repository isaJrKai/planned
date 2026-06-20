"""Generate the comprehensive audit document for the Planned Kids Savings app."""
import sys, os
from pathlib import Path

# Ensure docx is available
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    os.system("pip install python-docx --quiet")
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

OUT = Path("/home/z/my-project/download/Planned_App_Audit.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

# ---- Style setup ----
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
pf = style.paragraph_format
pf.line_spacing = 1.4
pf.space_after = Pt(4)

# Helper to add heading with custom color
def heading(text, level=1, color=(201, 168, 76)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def para(text, bold=False, italic=False, color=None, size=11, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table_row(table, cells):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

# ============================================================================
# COVER
# ============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("PLANNED")
title_run.font.size = Pt(36)
title_run.font.color.rgb = RGBColor(201, 168, 76)
title_run.bold = True
title.paragraph_format.space_before = Pt(80)
title.paragraph_format.space_after = Pt(0)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run("Family Savings & Allowance App")
sub_run.font.size = Pt(16)
sub_run.font.color.rgb = RGBColor(120, 120, 120)
sub_run.italic = True
subtitle.paragraph_format.space_after = Pt(40)

audtitle = doc.add_paragraph()
audtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = audtitle.add_run("Complete Application Audit")
ar.font.size = Pt(22)
ar.font.color.rgb = RGBColor(40, 40, 40)
ar.bold = True
audtitle.paragraph_format.space_after = Pt(8)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = date_p.add_run("Prepared June 2026 · v3.0")
dr.font.size = Pt(11)
dr.font.color.rgb = RGBColor(140, 140, 140)
dr.italic = True

# Divider
div = doc.add_paragraph()
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
div.add_run("◆ ◆ ◆").font.color.rgb = RGBColor(201, 168, 76)

doc.add_paragraph()  # spacer

# ============================================================================
# SECTION 1: EXECUTIVE SUMMARY
# ============================================================================
heading("1. Executive Summary", level=1)
para(
    "Planned is a Next.js 16 + TypeScript + Tailwind CSS + Zustand web application designed to help "
    "families teach children disciplined wealth-building through savings, investments, tokens, "
    "spending tracking, and now flexible goals. The app features a parent dashboard (QuickBooks-"
    "style organized with editorial Japandi-luxe aesthetic) and a per-child dashboard, plus a "
    "four-theme system (Onyx, Ivory, Blush, Oxblood), profile photo uploads, and a parent-set "
    "annual theme + monthly quote that appears at the bottom of every child dashboard."
)
para(
    "This audit covers the third major iteration of the app. All 14 bugs identified in the "
    "previous audit are confirmed fixed. The sidebar theme bug reported in this iteration has "
    "been resolved. Profile photos, the goals feature (with privacy + cadence + tabular graph "
    "view), and the family editorial footer are all new in v3.0."
)

# Summary stats table
para("Build summary at a glance:", bold=True, size=12)
stats_table = doc.add_table(rows=1, cols=2)
stats_table.style = "Light List Accent 1"
hdr = stats_table.rows[0].cells
hdr[0].text = "Metric"
hdr[1].text = "Value"
for r in hdr[0].paragraphs[0].runs: r.bold = True
for r in hdr[1].paragraphs[0].runs: r.bold = True
add_table_row(stats_table, ["Framework", "Next.js 16 (App Router) + TypeScript 5"])
add_table_row(stats_table, ["Styling", "Tailwind CSS 4 + custom editorial design system"])
add_table_row(stats_table, ["State", "Zustand (client-side, in-memory)"])
add_table_row(stats_table, ["Themes", "4 (Onyx / Ivory / Blush / Oxblood)"])
add_table_row(stats_table, ["Parent tabs", "7 (Overview / Children / Transactions / Investments / Tokens / Goals / Settings)"])
add_table_row(stats_table, ["Child tabs", "4 (Overview / Spending / Worksheet / Investments)"])
add_table_row(stats_table, ["Total components", "12+ (modals, charts, avatar, goals, theme switcher, footer, etc.)"])
add_table_row(stats_table, ["Original bugs fixed", "14 / 14 (100%)"])
add_table_row(stats_table, ["v3.0 bugs fixed", "1 / 1 (sidebar theme)"])
add_table_row(stats_table, ["ESLint errors", "0"])
add_table_row(stats_table, ["Console errors (browser)", "0"])

doc.add_page_break()

# ============================================================================
# SECTION 2: WHAT THE APP CAN DO
# ============================================================================
heading("2. What the App Can Do", level=1)

heading("2.1 Parent Dashboard", level=2)
bullet("View family-wide KPIs: total savings, total investments, this-month saved, circulating tokens", "Overview — ")
bullet("See 4 editorial SVG charts: 6-month savings trend line, per-child distribution donut, saved-vs-spent bars, per-child goal radial progress", "Visualizations — ")
bullet("Browse all children in a sortable table with avatars, current savings, goals, progress bars", "Children tab — ")
bullet("Award tokens to any child (parent buys at UGX 50/◈, child redeems at UGX 80/◈ — 60% incentive spread)", "Tokens — ")
bullet("View full transaction ledger with type pills, cash + token columns, child avatars", "Transactions — ")
bullet("Track all family investment positions with principal, current value, gain %, status", "Investments — ")
bullet("Create, edit, delete goals. Filter by owner + type. Tabular view with graph progress column. Per-owner progress bars.", "Goals — ")
bullet("Upload family profile photos, edit names, set annual theme + monthly quote, configure token economics, view family net worth", "Settings — ")

heading("2.2 Child Dashboard", level=2)
bullet("Hero balance with animated gold foil, goal progress, four stat cards (this month saved, total investments, token balance, linked account)", "Overview — ")
bullet("Per-child category breakdown with budget bars, spending log table", "Spending — ")
bullet("Monthly cash flow summary, net position, daily save average, goal progress", "Worksheet — ")
bullet("Portfolio summary, holdings table with gain %, click-through to detail modal", "Investments — ")
bullet("Save Money (with balance validation), Redeem Tokens, Log Spending, View Tokens", "Quick actions — ")
bullet("Annual theme banner + monthly quote at the bottom, signed 'Set with love · by Parent'", "Family footer — ")

heading("2.3 Goals Feature (New in v3.0)", level=2)
bullet("Mother, father, or any child can own goals", "Owner flexibility — ")
bullet("'Save' (build toward a target) or 'Spend Less' (stay under a budget cap)", "Goal types — ")
bullet("Weekly, monthly, or annual — period auto-resets via resetGoalPeriod()", "Cadence — ")
bullet("Private (owner only) or Revealed (whole family sees it) — enforced via viewerId filter", "Privacy — ")
bullet("Maximum 15 goals per owner (enforced in store.addGoal, returns error if exceeded)", "Cap — ")
bullet("Tabular with graph progress column (animated SVG bar + percentage + period label), plus per-owner progress graph", "Display — ")
bullet("Edit, delete, contribute to savings goals; confirm-delete modal prevents accidents", "Management — ")

heading("2.4 Profile Photos (New in v3.0)", level=2)
bullet("Avatar component with halo glow, falls back to initials on avatarColor when no photo", "Component — ")
bullet("File input → canvas resize to 256×256 → base64 JPEG stored in Zustand", "Pipeline — ")
bullet("Photos appear in parent topbar, children table, children cards, transactions table, tokens table, child dashboard header, goals table", "Surfaces — ")
bullet("Editable from Settings → Family Profiles card with name field + photo upload + remove", "Editor — ")

heading("2.5 Theme System (Reworked in v3.0)", level=2)
bullet("Deep green-black + bright gold shimmer", "Onyx — ")
bullet("'Hermès Private Bank' — warm aged-paper ivory + antique gold + olive/bronze chart palette", "Ivory — ")
bullet("'Chanel Boutique' — warmer blush + deep wine text + rose-gold primary + wine → rose-gold → gold foil", "Blush — ")
bullet("'Venetian Royal Chamber' — deeper oxblood + cream text + antique gold with bright pop + velvet surfaces", "Oxblood — ")
bullet("Each theme has its own --gold-foil-1/2/3 tokens so foil text is always visible against the background (was the root cause of 'themes not working' in v2.0)", "Theme-specific gold foil — ")
bullet("Theme-specific --card-shadow + --card-shadow-hover, .card-hover utility for lift effect", "Card depth — ")
bullet("Two radial gold glows slowly drift across the screen (32s + 40s animations) using theme-specific glow colors", "Ambient backdrop — ")
bullet("Persists to localStorage, loaded before paint via inline script (no FOUC)", "Persistence — ")

heading("2.6 Family Editorial (Added in v2.0)", level=2)
bullet("Parent sets a short motto (80 char limit) — appears as gold-foil banner at bottom of child dashboard", "Annual theme — ")
bullet("Parent sets a monthly quote (160 char limit) — appears in italic Playfair with ornamental rules + 'Set with love · by Parent'", "Monthly quote — ")
bullet("5 suggested annual themes + 8 suggested monthly quotes as one-click chips in Settings", "Suggestions — ")
bullet("Live preview card in Settings shows exactly what children will see", "Preview — ")

doc.add_page_break()

# ============================================================================
# SECTION 3: WHAT THE APP CAN'T DO (YET)
# ============================================================================
heading("3. What the App Can't Do (Yet)", level=1)

heading("3.1 Persistence Limitations", level=2)
bullet("All data lives in a Zustand store in browser memory. Refreshing the page resets everything to seed data. No Prisma/SQLite backend is connected yet, even though both are installed.")
bullet("Profile photos are stored as base64 strings in the Zustand store — they disappear on refresh. For production, these should be uploaded to object storage (S3, OSS, etc.).")
bullet("Annual theme + monthly quote do not survive a page refresh for the same reason.")

heading("3.2 Authentication & Multi-User", level=2)
bullet("There is no login system. Anyone who opens the app sees the same family data. NextAuth.js is installed but not configured.")
bullet("There is no concept of 'current user' — the privacy filter on goals takes a viewerId prop, but nothing actually sets it based on who's logged in.")
bullet("There is no parent vs child mode toggle — a child opening the app sees the parent dashboard (and could award themselves tokens).")

heading("3.3 Real Broker Integration", level=2)
bullet("The InvestmentDetailModal 'Invest More' button is honestly disabled with '(Coming Soon)' — no broker is connected. Real investment execution would require a securities API (e.g.elynker, Taproot, or local Ugandan broker integration).")
bullet("Investment current values are static mock numbers. No market data feed updates them.")

heading("3.4 Notifications & Scheduling", level=2)
bullet("No push notifications, email reminders, or SMS alerts. The notification bell in the topbar is purely decorative.")
bullet("No cron/scheduler for auto-streak logic, automatic period resets, or recurring allowance payouts. The resetGoalPeriod() method exists but must be called manually.")
bullet("No 'auto-award tokens for chores' automation.")

heading("3.5 Mobile App / PWA", level=2)
bullet("The web app is responsive but is not packaged as a PWA. No service worker, no installable manifest, no offline mode.")
bullet("No native iOS/Android apps. Children would need to access via browser.")

heading("3.6 Multi-Currency / Localization", level=2)
bullet("Hardcoded to UGX (Ugandan Shillings). No currency switcher. No i18n framework — UI is English-only despite next-intl being installed.")

heading("3.7 Reporting & Export", level=2)
bullet("No PDF statement generation. No CSV/Excel export of transactions. No printable monthly report.")
bullet("No email summaries for parents ('Your child saved UGX X this month').")

heading("3.8 Audit Trail / Undo", level=2)
bullet("No undo for any action. If a parent accidentally deletes a goal, it's gone.")
bullet("No audit log of who did what (would matter more once multi-user auth exists).")
bullet("No soft-delete — deleteGoal() removes immediately.")

doc.add_page_break()

# ============================================================================
# SECTION 4: WHAT THE APP HAS
# ============================================================================
heading("4. What the App Has", level=1)

heading("4.1 Data Model", level=2)
bullet("Child (id, name, age, avatarColor, avatarPhoto, currentAmount, goalAmount, goalName)")
bullet("ParentProfile (id, name, role, avatarColor, avatarPhoto)")
bullet("Account (linked external account per child with balance)")
bullet("Transaction (save / withdraw / invest / redeem / parent_give)")
bullet("SpendingEntry (per-child, per-category, with note + timestamp)")
bullet("SpendingCategory (6 seeded: Snacks, Airtime, School, Toys, Gifts, Transport)")
bullet("Investment (Equity / Bond / Savings Bond / Treasury Bill / Unit Trust)")
bullet("TokenLedgerEntry (parent_give / redeem)")
bullet("Goal (with ownerId, ownerKind, type, cadence, visibility, target, current, periodStart)")

heading("4.2 Components", level=2)
bullet("Avatar (with photo upload + initials fallback + halo glow)")
bullet("SaveMoneyModal, GiveTokensModal, AddSpendingModal, RedeemTokensModal, InvestmentDetailModal, GoalModal, ContributeModal, ConfirmDeleteModal")
bullet("ThemeSwitcher (4-swatch dropdown, persists to localStorage)")
bullet("ParentQuoteEditor (with live preview + suggestions)")
bullet("FamilyThemeFooter (annual theme banner + monthly quote)")
bullet("GoalsTab + GoalRow + ProfileEditorCard")
bullet("4 editorial SVG charts: SavingsTrendChart, DistributionDonut, CashFlowBars, GoalRadials")
bullet("ChildDashboard (4 tabs + floating action buttons)")

heading("4.3 Design System", level=2)
bullet("4 fully-tokenized themes via [data-theme] CSS variables")
bullet("Playfair Display serif + Inter sans + JetBrains Mono")
bullet("Custom CSS classes: surface-wood, surface-wood-strong, surface-flat, divider-gold, ornamental-rule, family-banner, halo-glow, micro-label, btn-gold, btn-outline, btn-ghost, input-editorial, ledger-table, progress-thin, pill-*")
bullet("Theme-specific gold foil gradients (--gold-foil-1/2/3 tokens)")
bullet("Theme-specific card shadows + ambient backdrop animation")

heading("4.4 Money Logic (Verified Correct)", level=2)
bullet("Save Money updates child.currentAmount AND debits linked account (was bug #1)")
bullet("thisMonthSaved filters by childId (was bug #2)")
bullet("totalInvested sums currentValue (was bug #3)")
bullet("parentTokenBalance subtracts redeemed tokens (was bug #4)")
bullet("Save modal validates against linked account balance (was bug #5)")
bullet("Investment 'Invest More' honestly disabled (was bug #6)")
bullet("Per-child spending categories computed live (was bug #7)")
bullet("Overview 'This Month' shows savings credits not spending (was bug #8)")
bullet("Shared date/budget helpers (was bug #9)")
bullet("getEncouragement randomized per call (was bug #10)")
bullet("Worksheet totalReceived filters to current month (was bug #11)")
bullet("addSpendingEntry surfaces immediately (was bug #12)")
bullet("No dead Prisma import (was bug #13)")
bullet("TOKEN_REDEEM_RATE constant used everywhere (was bug #14)")

doc.add_page_break()

# ============================================================================
# SECTION 5: WHAT THE APP LACKS
# ============================================================================
heading("5. What the App Lacks", level=1)

heading("5.1 Critical Gaps", level=2)
bullet("Real database — the #1 gap. Prisma + SQLite are installed but unused. All data resets on refresh.")
bullet("Authentication — no login, no current-user concept, no role-based access (parent vs child).")
bullet("Audit trail — no record of who did what, when. Critical for a money app.")

heading("5.2 Important Gaps", level=2)
bullet("Photo storage — base64 in memory is fine for MVP but won't scale. Needs object storage.")
bullet("Notifications — no way to nudge children or alert parents of milestones.")
bullet("Recurring automation — no scheduler for allowance payouts, period resets, streak awards.")
bullet("Real broker — investments are mock. Real execution needs a securities API.")
bullet("Multi-currency — hardcoded to UGX. Should support USD, KES, RWF, TZS at minimum.")
bullet("Reporting — no PDF/CSV export, no email summaries.")

heading("5.3 Nice-to-Have Gaps", level=2)
bullet("PWA packaging for installable mobile use")
bullet("Native iOS/Android apps")
bullet("Multi-language support (Luganda, Swahili, French)")
bullet("Family chat / comments on transactions")
bullet("Educational content module (financial literacy lessons)")
bullet("Gamification badges beyond tokens (e.g. 'Saver of the Month')")

doc.add_page_break()

# ============================================================================
# SECTION 6: WHAT CAN BE FIXED (vs WHAT CAN'T)
# ============================================================================
heading("6. What Can Be Fixed vs What Can't", level=1)

heading("6.1 Easily Fixable (Tech Debt)", level=2)
easy_table = doc.add_table(rows=1, cols=3)
easy_table.style = "Light List Accent 1"
h = easy_table.rows[0].cells
h[0].text = "Item"
h[1].text = "Effort"
h[2].text = "Notes"
for c in h:
    for r in c.paragraphs[0].runs: r.bold = True
add_table_row(easy_table, ["Connect Prisma + SQLite", "Medium", "Schema is ready in package.json. Replace Zustand seed with db queries."])
add_table_row(easy_table, ["Wire NextAuth.js", "Medium", "Installed. Add CredentialsProvider + role field on User model."])
add_table_row(easy_table, ["Move photos to S3/OSS", "Small", "Replace base64 in store with upload + URL field on Child/Parent."])
add_table_row(easy_table, ["Add CSV export", "Small", "Use react-csv or server-side stream."])
add_table_row(easy_table, ["Add PDF monthly statement", "Medium", "Use pdf skill with editorial template."])
add_table_row(easy_table, ["Add cron for period resets", "Small", "Vercel cron + API route calling resetGoalPeriod for all goals."])
add_table_row(easy_table, ["Add multi-currency", "Medium", "Use Intl.NumberFormat + currency field on family."])
add_table_row(easy_table, ["Add undo for delete goal", "Small", "Soft-delete + 5-second undo toast."])
add_table_row(easy_table, ["Add PWA manifest", "Small", "next-pwa package + manifest.json."])
add_table_row(easy_table, ["Prune unused dependencies", "Trivial", "next-auth, recharts, react-query, @dnd-kit, sharp, etc. — 40+ unused."])

heading("6.2 Hard but Possible", level=2)
add_table_row(easy_table, ["Real broker integration", "Large", "Need Ugandan securities API or partner. Compliance + custody issues."])
add_table_row(easy_table, ["Mobile native apps", "Large", "Reuse API + rebuild UI in React Native / Flutter."])
add_table_row(easy_table, ["Email notification system", "Medium", "Resend / SendGrid + templates + scheduler."])
add_table_row(easy_table, ["Audit trail with diffs", "Medium", "Event-sourcing pattern or per-table history columns."])

heading("6.3 Cannot Be Fixed (By Design)", level=2)
bullet("The app is client-side only — there's no server to 'remember' state. This is a design choice for the MVP, not a bug. The fix is to add a backend, which is item #1 in 6.1.")
bullet("The 'Invest More' button is disabled because there's no broker. This is HONEST behavior, not a bug — pretending to invest real money would be fraud. The fix is real broker integration, which is item #1 in 6.2.")
bullet("Children can see the parent dashboard because there's no auth. This is a design choice for the MVP. The fix is item #2 in 6.1.")

doc.add_page_break()

# ============================================================================
# SECTION 7: WHAT IS FINISHED
# ============================================================================
heading("7. What Is Finished (Production-Ready)", level=1)
bullet("All 14 money-logic bugs from v1.0 audit — fixed and browser-verified")
bullet("Sidebar theme bug from v3.0 — fixed (sidebar now uses bg-sidebar token, switches with theme)")
bullet("4-theme system with distinct identities + theme-specific gold foil + card shadows + ambient backdrop")
bullet("Parent dashboard with 7 tabs (Overview, Children, Transactions, Investments, Tokens, Goals, Settings)")
bullet("Child dashboard with 4 tabs (Overview, Spending, Worksheet, Investments) + family footer")
bullet("Goals feature: create/edit/delete/contribute, privacy filter, cadence, tabular view with graph progress, per-owner summary")
bullet("Profile photo upload with canvas resize, all avatars swapped to use Avatar component")
bullet("Annual theme + monthly quote with live preview + suggestions library")
bullet("4 editorial SVG charts on parent overview")
bullet("Token economy with 60% incentive spread")
bullet("All money flows: Save, Withdraw, Invest, Redeem, Award — verified end-to-end")
bullet("ESLint clean, no console errors, no page errors, responsive on mobile + desktop")

# ============================================================================
# SECTION 8: WHAT IS UNFINISHED
# ============================================================================
heading("8. What Is Unfinished (Known)", level=1)
bullet("No persistence — refresh = data loss (see 6.1 #1)")
bullet("No auth — anyone can do anything (see 6.1 #2)")
bullet("No real investments — 'Invest More' disabled (see 6.3)")
bullet("No notifications — bell is decorative (see 5.2)")
bullet("No reporting/export — data is trapped in the UI (see 5.2)")
bullet("No multi-currency — UGX only (see 5.2)")
bullet("40+ unused npm dependencies — bloat but no runtime impact (see 6.1 #10)")
bullet("Prisma schema exists but is unused — ready to wire up when backend is added")

# ============================================================================
# SECTION 9: v3.0 CHANGES (THIS ITERATION)
# ============================================================================
heading("9. v3.0 Changes (This Iteration)", level=1)

heading("9.1 Bug Fixes", level=2)
bullet("Sidebar was using hardcoded bg-[#0B0F0D] — replaced with bg-sidebar + var(--sidebar-border). Now switches correctly across all 4 themes (verified via getComputedStyle: dark=rgb(11,15,13), light=rgb(237,227,200), pink=rgb(245,221,216), red=rgb(17,3,3)).")
bullet("Hardcoded color in modals.tsx option element (bg-[#0E1310]) — replaced with bg-sidebar.")

heading("9.2 New Features", level=2)
bullet("Goals feature with 5 seeded goals (Emergency Fund, Family Holiday, Spend Less on Coffee, Zara's tablet, Enoch's weekly save). Parents and children can own goals. Privacy + cadence + cap of 15 per owner.")
bullet("Profile photo upload via Avatar component with canvas resize to 256×256, base64 JPEG stored in Zustand. All avatars throughout the app now use the Avatar component (topbar, tables, cards, child dashboard header, goals table).")
bullet("ParentProfile type + 2 seeded parents (Mama, Papa) with their own photos + names.")
bullet("Settings → Family Profiles card with avatar upload + name editing for all 5 family members.")
bullet("Goals nav item added between Tokens and Settings.")

heading("9.3 Verification", level=2)
bullet("Switched through all 4 themes — sidebar bg color confirmed to change in each (via getComputedStyle).")
bullet("Created a new 'Spend Less on Eating Out' goal with Monthly cadence + Private visibility — appeared in table immediately.")
bullet("Renamed Zara to 'Zara Namutebi-Okello' in Settings — propagated to children table on Overview tab.")
bullet("ESLint clean (0 errors, 0 warnings).")
bullet("No console errors, no page errors.")
bullet("Screenshots: 08-sidebar-theme-fix.png, 09-goals-tab.png, 10-settings-profiles.png")

doc.add_page_break()

# ============================================================================
# SECTION 10: RECOMMENDED NEXT STEPS
# ============================================================================
heading("10. Recommended Next Steps (Priority Order)", level=1)

heading("Priority 1 — Make It Real", level=2)
bullet("Connect Prisma + SQLite. Replace Zustand seed arrays with db queries. Add a /api/seed endpoint to populate demo data on first run.", "1.")
bullet("Wire NextAuth.js with CredentialsProvider. Add User model with role (parent/child). Add a login page. Hide parent dashboard behind parent role.", "2.")
bullet("Move photo uploads to S3/OSS. Add a /api/upload endpoint that returns a URL. Store URL on Child.avatarPhoto + ParentProfile.avatarPhoto.", "3.")

heading("Priority 2 — Make It Useful", level=2)
bullet("Add Vercel cron for automatic goal period resets (weekly/monthly/annual).", "4.")
bullet("Add CSV export button on Transactions tab. Use react-csv client-side.", "5.")
bullet("Add PDF monthly statement generator using the pdf skill.", "6.")
bullet("Add email notifications via Resend for: child saved X, child reached goal, parent awarded tokens, monthly summary.", "7.")

heading("Priority 3 — Make It Scale", level=2)
bullet("Add multi-currency support. Add a currency field on Family. Use Intl.NumberFormat everywhere.", "8.")
bullet("Add PWA manifest + service worker for installable mobile use.", "9.")
bullet("Add audit trail table (userId, action, entity, entityId, before, after, timestamp).", "10.")
bullet("Prune unused npm dependencies (40+ packages: next-auth is used now, but recharts, react-query, @dnd-kit, sharp, etc. are still unused).", "11.")

heading("Priority 4 — Make It Special", level=2)
bullet("Real broker integration (large effort, requires partner).", "12.")
bullet("Educational content module (financial literacy lessons for kids).", "13.")
bullet("Gamification badges beyond tokens (Saver of the Month, Streak Master, etc.).", "14.")
bullet("Family chat / comments on transactions.", "15.")

# Save
doc.save(str(OUT))
print(f"Audit saved to: {OUT}")
print(f"Size: {OUT.stat().st_size:,} bytes")
