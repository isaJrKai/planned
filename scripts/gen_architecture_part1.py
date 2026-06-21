"""Generate the Planned Production Architecture DOCX — Part 1: Setup + Sections 1-5."""
import sys, os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path("/home/z/my-project/public/Planned_Production_Architecture.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Styles
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35

GOLD = RGBColor(201, 168, 76)
DARK = RGBColor(30, 30, 30)
GRAY = RGBColor(120, 120, 120)
CODE_COLOR = RGBColor(40, 80, 40)

# ---- Helpers ----

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GOLD if level == 1 else DARK
    return h

def add_para(text, bold=False, italic=False, color=None, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    r.font.size = Pt(size)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_code(code_text, language=""):
    """Add a code block with monospace font + light shading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.line_spacing = 1.15
    # Add shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F0")
    pPr.append(shd)
    # Add border
    pBdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "CCCCCC")
        pBdr.append(border)
    pPr.append(pBdr)
    r = p.add_run(code_text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = CODE_COLOR
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def page_break():
    doc.add_page_break()

# ============================================================================
# COVER
# ============================================================================
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("PLANNED")
r.font.size = Pt(42)
r.font.color.rgb = GOLD
r.bold = True
cover.paragraph_format.space_before = Pt(100)
cover.paragraph_format.space_after = Pt(0)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Production Architecture Specification")
r.font.size = Pt(20)
r.font.color.rgb = DARK
r.bold = True
sub.paragraph_format.space_after = Pt(8)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Family Savings & Allowance Platform")
r.font.size = Pt(13)
r.font.color.rgb = GRAY
r.italic = True
sub2.paragraph_format.space_after = Pt(60)

# Meta
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Version 1.0  |  June 2026  |  Staff/Principal Engineering")
r.font.size = Pt(10)
r.font.color.rgb = GRAY
meta.paragraph_format.space_after = Pt(80)

# Scope box
scope = doc.add_paragraph()
scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = scope.add_run(
    "This document specifies the production infrastructure layer for the Planned\n"
    "family finance platform — designed to support 10,000+ families, 50,000+\n"
    "children, and millions of transactions."
)
r.font.size = Pt(11)
r.font.color.rgb = DARK
r.italic = True

page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
add_heading("Contents", level=1)
toc_items = [
    "1. System Architecture Diagram",
    "2. Folder Structure",
    "3. Database Schema",
    "4. Authentication Architecture",
    "5. Authorization Architecture",
    "6. API Architecture",
    "7. State Management Architecture",
    "8. Services Layer",
    "9. Repository Layer",
    "10. Notifications System",
    "11. Scheduler System",
    "12. Reporting System",
    "13. File Upload System",
    "14. Audit Trail System",
    "15. Security Architecture",
    "16. Deployment Architecture",
    "17. Environment Variables",
    "18. Migration Plan",
    "19. Scaling Strategy",
    "20. Implementation Roadmap",
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(item)
    r.font.size = Pt(11)

page_break()

# ============================================================================
# SECTION 1: SYSTEM ARCHITECTURE DIAGRAM
# ============================================================================
add_heading("1. System Architecture Diagram", level=1)

add_para("Overview", bold=True, size=12)
add_para(
    "Planned is a multi-tenant family finance platform built on Next.js 16 with a "
    "clear separation between the presentation layer (React components), the "
    "application layer (server actions + route handlers), the domain layer "
    "(services + repositories), and the persistence layer (PostgreSQL + object "
    "storage). The architecture follows Clean Architecture + DDD principles so "
    "that business rules are isolated from infrastructure concerns."
)

add_para("High-Level Architecture (Text Diagram)", bold=True, size=12)
add_code("""┌─────────────────────────────────────────────────────────────────────┐
│                        CLIENT (Browser / PWA)                        │
│  React Components · Zustand (UI state only) · Tailwind CSS           │
└──────────────────────────────┬──────────────────────────────────────┘
                               │ HTTPS
┌──────────────────────────────▼──────────────────────────────────────┐
│                    NEXT.JS 16 (VERCEL EDGE)                          │
│  ┌─────────────┐  ┌──────────────┐  ┌────────────────────────────┐  │
│  │  Middleware  │  │  Server      │  │  Route Handlers /api/*     │  │
│  │  (auth,      │  │  Actions     │  │  (REST endpoints)          │  │
│  │   CSRF,      │  │  (mutations) │  │                            │  │
│  │   rate-limit)│  │              │  │                            │  │
│  └──────┬───────┘  └──────┬───────┘  └────────────┬───────────────┘  │
│         │                 │                       │                   │
│  ┌──────▼─────────────────▼───────────────────────▼───────────────┐  │
│  │                     APPLICATION LAYER                          │  │
│  │  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐         │  │
│  │  │ Auth     │ │ Token    │ │ Goal     │ │ Report   │         │  │
│  │  │ Service  │ │ Service  │ │ Service  │ │ Service  │         │  │
│  │  │ (RBAC)   │ │ (black-  │ │          │ │          │         │  │
│  │  │          │ │  box)    │ │          │ │          │         │  │
│  │  └────┬─────┘ └────┬─────┘ └────┬─────┘ └────┬─────┘         │  │
│  └───────┼────────────┼────────────┼────────────┼────────────────┘  │
│          │            │            │            │                    │
│  ┌───────▼────────────▼────────────▼────────────▼────────────────┐  │
│  │                    REPOSITORY LAYER                            │  │
│  │  FamilyRepo · ChildRepo · TransactionRepo · GoalRepo ·        │  │
│  │  InvestmentRepo · SpendingRepo · TokenLedgerRepo · AuditRepo  │  │
│  └───────────────────────────────┬───────────────────────────────┘  │
│                                  │                                   │
│  ┌───────────────────────────────▼───────────────────────────────┐  │
│  │                    INFRASTRUCTURE LAYER                        │  │
│  │  Prisma ORM · Zod Validators · Logger · Cache · Queue         │  │
│  └───────────────────────────────┬───────────────────────────────┘  │
└──────────────────────────────────┼──────────────────────────────────┘
                                   │
          ┌────────────────────────┼────────────────────────┐
          │                        │                        │
   ┌──────▼──────┐          ┌──────▼──────┐         ┌──────▼──────┐
   │ PostgreSQL  │          │    S3 /     │         │   Redis     │
   │ (primary)   │          │  Cloudinary │         │  (cache +   │
   │             │          │  (avatars)  │         │   queue)    │
   └─────────────┘          └─────────────┘         └─────────────┘
          │                        │                        │
   ┌──────▼──────┐          ┌──────▼──────┐         ┌──────▼──────┐
   │  PgBouncer  │          │   CDN       │         │  BullMQ     │
   │  (pooling)  │          │  (delivery) │         │  Workers    │
   └─────────────┘          └─────────────┘         │  (jobs)     │
                                                    └─────────────┘

EXTERNAL SERVICES:
  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐
  │  Resend      │  │  Vercel Cron │  │  Sentry      │  │  Stripe      │
  │  (email)     │  │  (scheduler) │  │  (errors)    │  │  (future)    │
  └──────────────┘  └──────────────┘  └──────────────┘  └──────────────┘""")

add_para("Reasoning", bold=True, size=12)
add_bullet("The presentation layer contains ZERO business logic — components only render state and dispatch intents.", "Separation of concerns: ")
add_bullet("Server Actions handle mutations (type-safe, no manual fetch); Route Handlers handle queries + webhooks + file uploads.", "Dual API surface: ")
add_bullet("Services own business rules (TokenService is a black-box interface); Repositories own data access; Infrastructure owns cross-cutting concerns.", "Layered isolation: ")
add_bullet("PostgreSQL for transactional data, S3/Cloudinary for blobs, Redis for cache + job queue. Each store is best-in-class for its workload.", "Polyglot persistence: ")
add_bullet("Vercel handles auto-scaling + edge caching. Background jobs run on Vercel Functions triggered by Vercel Cron, with BullMQ + Redis for heavy processing.", "Serverless-first: ")

add_para("Tradeoffs", bold=True, size=12)
add_bullet("Pro: simpler deployment, no container orchestration. Con: 60s function timeout on Vercel — long-running reports must be queued.", "Serverless vs. containers: ")
add_bullet("Pro: type-safe, ergonomic. Con: harder to cache than GET endpoints. Mitigation: use React Query for client-side caching of server-action reads.", "Server Actions vs. REST: ")
add_bullet("Pro: familiar, excellent DX. Con: 10K families × 5 members × 1000 transactions = 50M rows — needs partitioning by year + familyId. Handled in Section 19.", "PostgreSQL vs. DynamoDB: ")

page_break()

# ============================================================================
# SECTION 2: FOLDER STRUCTURE
# ============================================================================
add_heading("2. Folder Structure", level=1)

add_para("Production folder structure organized by feature + layer:", bold=True, size=12)
add_code("""planned/
├── prisma/
│   ├── schema.prisma              # Full production schema (Section 3)
│   ├── migrations/                # Versioned SQL migrations
│   └── seed.ts                    # Demo data seeder
│
├── src/
│   ├── app/                       # Next.js App Router
│   │   ├── (auth)/                # Auth route group (login, register, forgot)
│   │   │   ├── login/page.tsx
│   │   │   ├── register/page.tsx
│   │   │   └── forgot-password/page.tsx
│   │   ├── (dashboard)/           # Protected route group
│   │   │   ├── layout.tsx         # Auth guard + sidebar shell
│   │   │   ├── page.tsx           # Parent overview
│   │   │   ├── children/page.tsx
│   │   │   ├── goals/page.tsx
│   │   │   ├── transactions/page.tsx
│   │   │   ├── investments/page.tsx
│   │   │   ├── tokens/page.tsx
│   │   │   ├── reports/page.tsx
│   │   │   └── settings/page.tsx
│   │   ├── (child)/               # Child-scoped route group
│   │   │   ├── layout.tsx         # Child auth guard
│   │   │   └── [childId]/page.tsx
│   │   ├── api/                   # Route Handlers (REST)
│   │   │   ├── auth/[...nextauth]/route.ts
│   │   │   ├── families/route.ts
│   │   │   ├── children/route.ts
│   │   │   ├── goals/route.ts
│   │   │   ├── transactions/route.ts
│   │   │   ├── uploads/route.ts   # Presigned URL generation
│   │   │   ├── reports/route.ts
│   │   │   └── webhooks/          # External webhooks
│   │   │       └── resend/route.ts
│   │   ├── layout.tsx             # Root layout (fonts, providers)
│   │   └── globals.css            # 4-theme design system
│   │
│   ├── components/                # Shared UI components
│   │   ├── ui/                    # shadcn/ui primitives (button, card, etc.)
│   │   ├── charts/                # Editorial SVG charts
│   │   ├── modals/                # Shared modals
│   │   └── providers/             # React context providers
│   │       ├── auth-provider.tsx
│   │       ├── theme-provider.tsx
│   │       └── query-provider.tsx
│   │
│   ├── features/                  # Feature modules (DDD)
│   │   ├── auth/
│   │   │   ├── components/        # LoginForm, RegisterForm
│   │   │   ├── actions.ts         # Server actions: login, register, logout
│   │   │   ├── schemas.ts         # Zod validation schemas
│   │   │   └── types.ts
│   │   ├── family/
│   │   │   ├── components/
│   │   │   ├── actions.ts
│   │   │   ├── schemas.ts
│   │   │   └── types.ts
│   │   ├── children/
│   │   │   ├── components/
│   │   │   ├── actions.ts
│   │   │   ├── schemas.ts
│   │   │   └── types.ts
│   │   ├── transactions/
│   │   │   ├── components/
│   │   │   ├── actions.ts         # save, withdraw, invest, redeem
│   │   │   ├── schemas.ts
│   │   │   └── types.ts
│   │   ├── goals/
│   │   │   ├── components/
│   │   │   ├── actions.ts         # create, update, delete, contribute
│   │   │   ├── schemas.ts
│   │   │   └── types.ts
│   │   ├── investments/
│   │   │   ├── components/
│   │   │   ├── actions.ts
│   │   │   ├── schemas.ts
│   │   │   └── types.ts
│   │   ├── spending/
│   │   │   ├── components/
│   │   │   ├── actions.ts
│   │   │   ├── schemas.ts
│   │   │   └── types.ts
│   │   ├── tokens/                # Token economy (BLACK-BOX)
│   │   │   ├── components/        # UI only — no business logic
│   │   │   ├── actions.ts         # Calls TokenService interface
│   │   │   ├── types.ts
│   │   │   └── README.md          # "DO NOT modify token calculations"
│   │   ├── reports/
│   │   │   ├── components/
│   │   │   ├── actions.ts         # generate, download, email
│   │   │   └── types.ts
│   │   ├── notifications/
│   │   │   ├── components/        # Bell dropdown, toast
│   │   │   ├── actions.ts
│   │   │   └── types.ts
│   │   └── uploads/
│   │       ├── components/        # AvatarUploader
│   │       ├── actions.ts         # requestPresignedUrl, confirmUpload
│   │       └── types.ts
│   │
│   ├── lib/                       # Shared utilities (cross-feature)
│   │   ├── db.ts                  # Prisma client singleton
│   │   ├── auth.ts                # NextAuth configuration
│   │   ├── validations.ts         # Shared Zod schemas (email, password)
│   │   ├── errors.ts              # Typed error classes
│   │   ├── logger.ts              # Structured logger (pino)
│   │   ├── cache.ts               # Redis cache wrapper
│   │   ├── rate-limit.ts          # Rate limiter (Redis-backed)
│   │   ├── crypto.ts              # Password hashing, token generation
│   │   ├── currency.ts            # Multi-currency formatter
│   │   └── utils.ts               # cn(), formatters
│   │
│   ├── server/                    # Server-only code (never imported by client)
│   │   ├── repositories/          # Data access layer
│   │   │   ├── family.repository.ts
│   │   │   ├── child.repository.ts
│   │   │   ├── transaction.repository.ts
│   │   │   ├── goal.repository.ts
│   │   │   ├── investment.repository.ts
│   │   │   ├── spending.repository.ts
│   │   │   ├── token-ledger.repository.ts
│   │   │   ├── audit.repository.ts
│   │   │   └── base.repository.ts     # Generic CRUD base
│   │   ├── services/              # Business logic layer
│   │   │   ├── auth.service.ts
│   │   │   ├── token.service.ts       # INTERFACE ONLY (black-box)
│   │   │   ├── goal.service.ts
│   │   │   ├── transaction.service.ts
│   │   │   ├── investment.service.ts
│   │   │   ├── report.service.ts
│   │   │   ├── notification.service.ts
│   │   │   ├── upload.service.ts
│   │   │   ├── audit.service.ts
│   │   │   └── scheduler.service.ts
│   │   ├── permissions/           # RBAC authorization
│   │   │   ├── roles.ts           # ADMIN, PARENT, CHILD
│   │   │   ├── permissions.ts     # Permission matrix
│   │   │   └── guards.ts          # can(), requirePermission()
│   │   ├── jobs/                  # Background job definitions
│   │   │   ├── monthly-summary.job.ts
│   │   │   ├── goal-reset.job.ts
│   │   │   ├── streak-calc.job.ts
│   │   │   └── queue.ts           # BullMQ queue setup
│   │   └── emails/                # Email templates + sending
│   │       ├── templates/
│   │       │   ├── monthly-summary.tsx
│   │       │   ├── goal-achieved.tsx
│   │       │   └── password-reset.tsx
│   │       └── sender.ts
│   │
│   ├── hooks/                     # React hooks (client)
│   │   ├── use-auth.ts
│   │   ├── use-family.ts
│   │   ├── use-goals.ts
│   │   └── use-notifications.ts
│   │
│   ├── stores/                    # Zustand stores (UI state ONLY)
│   │   ├── ui.store.ts            # theme, sidebar, modals, filters
│   │   └── optimistic.store.ts    # Optimistic update coordination
│   │
│   ├── types/                     # Shared TypeScript types
│   │   ├── api.ts                 # Request/response types
│   │   ├── domain.ts              # Domain entity types
│   │   └── events.ts              # Domain event types
│   │
│   └── middleware.ts              # Auth + CSRF + rate-limit middleware
│
├── public/                        # Static assets
├── .env.example                   # Environment variable template
├── next.config.ts
├── tailwind.config.ts
├── tsconfig.json
├── package.json
└── README.md""")

add_para("Why each folder exists:", bold=True, size=12)
add_bullet("Isolates auth pages from the main app. Route groups don't affect URLs but enable distinct layouts.", "(auth) / (dashboard) / (child) route groups: ")
add_bullet("Each feature is a self-contained module with its own components, server actions, Zod schemas, and types. No feature imports from another feature — they communicate via server actions + events.", "features/: ")
add_bullet("Server-only code lives here. The `server/` prefix makes it visually obvious this code must never be imported by client components (Next.js will error if you try).", "server/: ")
add_bullet("One repository per aggregate root. Repositories return domain entities, not Prisma rows. This isolates the ORM — if we switch from Prisma to Drizzle, only repositories change.", "server/repositories/: ")
add_bullet("Services own business rules. TokenService is an interface only — the implementation is a black-box that must never be modified without domain-expert sign-off.", "server/services/: ")
add_bullet("UI state only. No domain data. The moment you put a `transactions` array in Zustand, you've broken the architecture — use React Query for server state.", "stores/: ")

page_break()

# ============================================================================
# SECTION 3: DATABASE SCHEMA
# ============================================================================
add_heading("3. Database Schema", level=1)

add_para("Production Prisma schema targeting PostgreSQL with UUID primary keys, soft deletes, optimistic concurrency, and full indexing:", bold=True, size=12)

add_code("""// prisma/schema.prisma

generator client {
  provider = "prisma-client-js"
}

datasource db {
  provider = "postgresql"
  url      = env("DATABASE_URL")
}

// ---- Enums ----------------------------------------------------------------

enum Role {
  ADMIN
  PARENT
  CHILD
}

enum TransactionType {
  SAVE
  WITHDRAW
  INVEST
  REDEEM
  PARENT_GIVE
}

enum GoalType {
  SAVE
  SPEND_LESS
}

enum GoalCadence {
  WEEKLY
  MONTHLY
  ANNUAL
}

enum GoalVisibility {
  PRIVATE
  REVEALED
}

enum InvestmentType {
  EQUITY
  BOND
  SAVINGS_BOND
  TREASURY_BILL
  UNIT_TRUST
}

enum InvestmentStatus {
  ACTIVE
  CLOSED
}

enum CurrencyCode {
  UGX
  USD
  KES
  TZS
  RWF
}

enum NotificationChannel {
  IN_APP
  EMAIL
  PUSH
}

enum NotificationStatus {
  UNREAD
  READ
  DISMISSED
}

enum AuditAction {
  LOGIN
  LOGOUT
  CREATE
  UPDATE
  DELETE
  TOKEN_AWARD
  TOKEN_REDEEM
  GOAL_CONTRIBUTE
  SETTINGS_CHANGE
  UPLOAD
}

// ---- Core Family Entities -------------------------------------------------

model Family {
  id          String   @id @default(uuid())
  name        String
  currency    CurrencyCode @default(UGX)
  annualTheme String   @default("")
  monthlyQuote String  @default("")
  createdAt   DateTime @default(now())
  updatedAt   DateTime @updatedAt
  deletedAt   DateTime?  // Soft delete

  users       User[]
  children    Child[]
  accounts    Account[]
  categories  SpendingCategory[]
  settings    FamilySettings?

  @@index([deletedAt])
}

model User {
  id              String   @id @default(uuid())
  email           String   @unique
  passwordHash    String?
  name            String
  role            Role     @default(PARENT)
  familyId        String
  emailVerified   DateTime?
  image           String?
  failedLoginAttempts Int  @default(0)
  lockedUntil     DateTime?
  createdAt       DateTime @default(now())
  updatedAt       DateTime @updatedAt
  deletedAt       DateTime?

  family          Family   @relation(fields: [familyId], references: [id], onDelete: Cascade)
  sessions        Session[]
  accounts        Account[]      // For social login linking
  auditLogs       AuditLog[]
  notifications   Notification[]
  devices         Device[]
  pushSubscriptions PushSubscription[]

  @@index([familyId])
  @@index([email])
  @@index([deletedAt])
}

model Session {
  id            String   @id @default(uuid())
  sessionToken  String   @unique
  userId        String
  expires       DateTime
  ipAddress     String?
  userAgent     String?
  createdAt     DateTime @default(now())

  user          User     @relation(fields: [userId], references: [id], onDelete: Cascade)

  @@index([userId])
  @@index([expires])
}

// ---- Family Members -------------------------------------------------------

model Child {
  id            String   @id @default(uuid())
  familyId      String
  name          String
  age           Int
  avatarColor   String
  avatarPhoto   String?
  currentAmount Int      @default(0)  // UGX live savings balance
  goalAmount    Int      @default(0)
  goalName      String
  pinHash       String?  // 4-digit PIN for child login
  createdAt     DateTime @default(now())
  updatedAt     DateTime @updatedAt
  deletedAt     DateTime?
  version       Int      @default(1)  // Optimistic concurrency

  family        Family   @relation(fields: [familyId], references: [id], onDelete: Cascade)
  accounts      Account[]
  transactions  Transaction[]
  spending      SpendingEntry[]
  investments   Investment[]
  tokenLedger   TokenLedgerEntry[]
  goals         Goal[]   @relation("ChildGoals")

  @@index([familyId])
  @@index([deletedAt])
}

model ParentProfile {
  id            String   @id @default(uuid())
  userId        String   @unique  // 1:1 with User (role=PARENT)
  familyId      String
  name          String
  role          String   // "Mother" | "Father" | "Guardian"
  avatarColor   String
  avatarPhoto   String?
  createdAt     DateTime @default(now())
  updatedAt     DateTime @updatedAt

  user          User     @relation(fields: [userId], references: [id], onDelete: Cascade)
  family        Family   @relation(fields: [familyId], references: [id])
  spending      SpendingEntry[]
  goals         Goal[]   @relation("ParentGoals")

  @@index([familyId])
}

// ---- Financial Entities ---------------------------------------------------

model Account {
  id        String   @id @default(uuid())
  familyId  String
  childId   String?
  name      String   // e.g. "Stanbic Junior"
  bankName  String?
  accountNumber String?  // Last 4 digits only, encrypted
  balance   Int      @default(0)
  currency  CurrencyCode @default(UGX)
  createdAt DateTime @default(now())
  updatedAt DateTime @updatedAt
  deletedAt DateTime?
  version   Int      @default(1)

  family    Family   @relation(fields: [familyId], references: [id], onDelete: Cascade)
  child     Child?   @relation(fields: [childId], references: [id], onDelete: Cascade)
  transactions Transaction[]

  @@index([familyId])
  @@index([childId])
}

model Transaction {
  id           String   @id @default(uuid())
  familyId     String
  childId      String
  type         TransactionType
  amount       Int      @default(0)
  tokenDelta   Int      @default(0)
  accountId    String?
  investmentId String?
  goalId       String?  // Optional: link to a goal
  note         String
  timestamp    DateTime @default(now())
  createdAt    DateTime @default(now())
  version      Int      @default(1)

  family       Family   @relation(fields: [familyId], references: [id], onDelete: Cascade)
  child        Child    @relation(fields: [childId], references: [id], onDelete: Cascade)
  account      Account? @relation(fields: [accountId], references: [id])
  investment   Investment? @relation(fields: [investmentId], references: [id])
  goal         Goal?    @relation(fields: [goalId], references: [id])
  auditLogs    AuditLog[]

  @@index([familyId, timestamp])
  @@index([childId, timestamp])
  @@index([type])
}

model SpendingCategory {
  id        String   @id @default(uuid())
  familyId  String
  name      String
  budget    Int      @default(0)
  createdAt DateTime @default(now())

  family    Family   @relation(fields: [familyId], references: [id], onDelete: Cascade)
  entries   SpendingEntry[]

  @@unique([familyId, name])
  @@index([familyId])
}

model SpendingEntry {
  id         String   @id @default(uuid())
  familyId   String
  ownerId    String   // childId OR parentProfileId
  ownerKind  String   // "parent" | "child"
  ownerName  String
  categoryId String?
  category   String   // Denormalized for historical records
  amount     Int      @default(0)
  note       String
  timestamp  DateTime @default(now())
  createdAt  DateTime @default(now())

  family     Family   @relation(fields: [familyId], references: [id], onDelete: Cascade)
  child      Child?   @relation(fields: [childId], references: [id], onDelete: Cascade)
  parent     ParentProfile? @relation(fields: [parentId], references: [id], onDelete: Cascade)
  categoryRef SpendingCategory? @relation(fields: [categoryId], references: [id])
  childId    String?
  parentId   String?

  @@index([familyId, timestamp])
  @@index([ownerId, timestamp])
}

model Investment {
  id             String   @id @default(uuid())
  familyId       String
  childId        String
  name           String
  type           InvestmentType
  amountInvested Int      @default(0)
  currentValue   Int      @default(0)
  status         InvestmentStatus @default(ACTIVE)
  openedAt       DateTime @default(now())
  closedAt       DateTime?
  createdAt      DateTime @default(now())
  updatedAt      DateTime @updatedAt
  version        Int      @default(1)

  family         Family   @relation(fields: [familyId], references: [id], onDelete: Cascade)
  child          Child    @relation(fields: [childId], references: [id], onDelete: Cascade)
  transactions   Transaction[]

  @@index([familyId, status])
  @@index([childId, status])
}

model TokenLedgerEntry {
  id        String   @id @default(uuid())
  familyId  String
  childId   String
  type      String   // "parent_give" | "redeem"
  tokens    Int      @default(0)
  note      String
  timestamp DateTime @default(now())
  createdAt DateTime @default(now())

  family    Family   @relation(fields: [familyId], references: [id], onDelete: Cascade)
  child     Child    @relation(fields: [childId], references: [id], onDelete: Cascade)

  @@index([familyId, childId, timestamp])
  @@index([childId, type])
}

// ---- Goals ----------------------------------------------------------------

model Goal {
  id            String   @id @default(uuid())
  familyId      String
  ownerId       String
  ownerKind     String   // "parent" | "child"
  ownerName     String
  title         String
  type          GoalType
  cadence       GoalCadence
  visibility    GoalVisibility
  targetAmount  Int      @default(0)
  currentAmount Int      @default(0)
  periodStart   DateTime @default(now())
  note          String?
  createdAt     DateTime @default(now())
  updatedAt     DateTime @updatedAt
  deletedAt     DateTime?
  version       Int      @default(1)

  family        Family   @relation(fields: [familyId], references: [id], onDelete: Cascade)
  child         Child?   @relation("ChildGoals", fields: [childId], references: [id], onDelete: Cascade)
  parent        ParentProfile? @relation("ParentGoals", fields: [parentId], references: [id], onDelete: Cascade)
  childId       String?
  parentId      String?
  contributions GoalContribution[]
  transactions  Transaction[]

  @@index([familyId, ownerId])
  @@index([ownerId, visibility])
  @@index([deletedAt])
}

model GoalContribution {
  id        String   @id @default(uuid())
  goalId    String
  amount    Int
  note      String?
  contributedBy String  // userId or childId
  contributedAt DateTime @default(now())
  createdAt DateTime @default(now())

  goal      Goal     @relation(fields: [goalId], references: [id], onDelete: Cascade)

  @@index([goalId, contributedAt])
}

// ---- Notifications --------------------------------------------------------

model Notification {
  id        String   @id @default(uuid())
  userId    String
  channel   NotificationChannel
  status    NotificationStatus @default(UNREAD)
  title     String
  body      String
  metadata  Json?
  createdAt DateTime @default(now())
  readAt    DateTime?

  user      User     @relation(fields: [userId], references: [id], onDelete: Cascade)

  @@index([userId, status])
  @@index([createdAt])
}

// ---- Audit Trail ----------------------------------------------------------

model AuditLog {
  id        String   @id @default(uuid())
  userId    String?
  familyId  String?
  action    AuditAction
  entityType String
  entityId  String
  before    Json?
  after     Json?
  ipAddress String?
  userAgent String?
  createdAt DateTime @default(now())

  user      User?    @relation(fields: [userId], references: [id])

  @@index([userId, createdAt])
  @@index([familyId, createdAt])
  @@index([entityType, entityId])
  @@index([action])
}

// ---- File Uploads ---------------------------------------------------------

model FileUpload {
  id          String   @id @default(uuid())
  familyId    String
  uploaderId  String
  filename    String
  mimeType    String
  size        Int
  storageKey  String   // S3 key or Cloudinary public_id
  storageUrl  String   // CDN URL
  thumbnailUrl String?
  createdAt   DateTime @default(now())
  deletedAt   DateTime?

  @@index([familyId, createdAt])
  @@index([storageKey])
}

// ---- Devices + Push -------------------------------------------------------

model Device {
  id        String   @id @default(uuid())
  userId    String
  name      String
  userAgent String?
  lastSeenAt DateTime @default(now())
  createdAt DateTime @default(now())

  user      User     @relation(fields: [userId], references: [id], onDelete: Cascade)

  @@index([userId])
}

model PushSubscription {
  id        String   @id @default(uuid())
  userId    String
  endpoint  String
  keys      Json
  createdAt DateTime @default(now())
  deactivatedAt DateTime?

  user      User     @relation(fields: [userId], references: [id], onDelete: Cascade)

  @@index([userId])
}

// ---- Email Log ------------------------------------------------------------

model EmailLog {
  id        String   @id @default(uuid())
  to        String
  subject   String
  template  String
  status    String   // "sent" | "failed" | "bounced"
  providerId String?  // Resend message ID
  error     String?
  createdAt DateTime @default(now())

  @@index([to, createdAt])
  @@index([status])
}

// ---- Family Settings (singleton per family) -------------------------------

model FamilySettings {
  id              String   @id @default(uuid())
  familyId        String   @unique
  notifyOnSave    Boolean  @default(true)
  notifyOnGoalReached Boolean @default(true)
  notifyOnTokenAward Boolean @default(true)
  monthlyReportDay Int    @default(1)  // Day of month to send report
  createdAt       DateTime @default(now())
  updatedAt       DateTime @updatedAt

  family          Family   @relation(fields: [familyId], references: [id], onDelete: Cascade)
}""", language="prisma")

add_para("ERD Explanation", bold=True, size=12)
add_bullet("top-level tenant boundary. All other entities have familyId for row-level security + multi-tenant isolation.", "Family: ")
add_bullet("email + password (PARENT) or email + password (CHILD with limited permissions). Children log in with a 4-digit PIN instead of password — simpler for kids.", "User: ")
add_bullet("all financial entities are scoped to familyId. Indexes on [familyId, timestamp] enable fast per-family queries even at 50M+ rows.", "Transaction/Spending/Investment: ")
add_bullet("version field on Child, Account, Investment, Goal enables optimistic concurrency — prevents lost updates when two parents edit simultaneously.", "Optimistic concurrency: ")
add_bullet("deletedAt on Family, User, Child, Account, Goal, FileUpload. Queries filter WHERE deletedAt IS NULL. A cron job hard-deletes after 90 days.", "Soft deletes: ")
add_bullet("GoalContribution tracks individual contributions to a goal (who, how much, when) — separate from the goal's currentAmount total.", "GoalContribution: ")

add_para("Migration Strategy", bold=True, size=12)
add_bullet("Use prisma migrate dev for local development. Each migration is a versioned SQL file checked into git.")
add_bullet("On deployment, prisma migrate deploy runs pending migrations in order. Never use db:push in production — it doesn't create migration files.")
add_bullet("For the SQLite-to-PostgreSQL transition (MVP → production), write a one-time data migration script that reads from SQLite and inserts into PostgreSQL via Prisma.")
add_bullet("For schema changes on large tables (e.g. adding a column to Transaction with 50M rows), use expand-then-contract: add the nullable column first, backfill in batches, then make it required in a follow-up migration.")

add_para("Seed Strategy", bold=True, size=12)
add_bullet("scripts/seed.ts creates 3 demo families with full data for local dev + staging.")
add_bullet("scripts/seed-production.ts creates only the ADMIN user + family on first deployment.")
add_bullet("Never seed production with demo data. Use a separate flag: SEED_DEMO_DATA=true only on local/staging.")

page_break()

# Save Part 1
doc.save(str(OUT))
print(f"Part 1 saved: {OUT}")
print(f"Size so far: {OUT.stat().st_size:,} bytes")
