"""Part 3: Sections 10-20 (Notifications, Scheduler, Reporting, Uploads, Audit, Security, Deployment, Env, Migration, Scaling, Roadmap)."""
import sys, os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path("/home/z/my-project/public/Planned_Production_Architecture.docx")
doc = Document(str(OUT))

GOLD = RGBColor(201, 168, 76)
DARK = RGBColor(30, 30, 30)
GRAY = RGBColor(120, 120, 120)
CODE_COLOR = RGBColor(40, 80, 40)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GOLD if level == 1 else DARK
    return h

def add_para(text, bold=False, italic=False, color=None, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    r.font.size = Pt(size)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_code(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.line_spacing = 1.15
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F0")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "CCCCCC")
        pBdr.append(border)
    pPr.append(pBdr)
    r = p.add_run(code_text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = CODE_COLOR
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def page_break():
    doc.add_page_break()

# ============================================================================
# SECTION 10: NOTIFICATIONS SYSTEM
# ============================================================================
page_break()
add_heading("10. Notifications System", level=1)

add_para("Multi-channel notifications (in-app, email, push) with preferences + retry logic.", bold=True, size=12)

add_code("""// src/server/services/notification.service.ts
import { db } from "@/lib/db";
import { emailSender } from "@/server/emails/sender";
import { pushSender } from "@/server/notifications/push-sender";
import { logger } from "@/lib/logger";

type NotificationEvent =
  | { type: "GOAL_REACHED"; goalId: string; ownerId: string }
  | { type: "SAVINGS_MILESTONE"; childId: string; amount: number }
  | { type: "TOKEN_AWARDED"; childId: string; tokens: number; awardedBy: string }
  | { type: "MONTHLY_SUMMARY"; familyId: string }
  | { type: "INVESTMENT_UPDATE"; childId: string; investmentId: string };

export const NotificationService = {
  async notify(event: NotificationEvent) {
    const recipients = await this.resolveRecipients(event);

    for (const recipient of recipients) {
      const prefs = await this.getPreferences(recipient.id);

      // Always create in-app notification
      await db.notification.create({
        data: {
          userId: recipient.id,
          channel: "IN_APP",
          title: this.getTitle(event),
          body: this.getBody(event),
          metadata: event as any,
        },
      });

      // Send email if enabled
      if (prefs.email && this.shouldEmail(event)) {
        await this.queueEmail(recipient, event);
      }

      // Send push if enabled
      if (prefs.push) {
        await this.queuePush(recipient, event);
      }
    }
  },

  async notifyGoalReached(goal: any) {
    await this.notify({ type: "GOAL_REACHED", goalId: goal.id, ownerId: goal.ownerId });
  },

  async notifyTokenAwarded(childId: string, tokens: number, awardedBy: string) {
    await this.notify({ type: "TOKEN_AWARDED", childId, tokens, awardedBy });
  },

  private async resolveRecipients(event: NotificationEvent): Promise<User[]> {
    switch (event.type) {
      case "GOAL_REACHED":
        // Notify goal owner + all parents in family
        const goal = await db.goal.findUnique({ where: { id: event.goalId } });
        const parents = await db.user.findMany({
          where: { familyId: goal.familyId, role: "PARENT" },
        });
        return parents;

      case "TOKEN_AWARDED":
        // Notify the child + their parents
        const child = await db.child.findUnique({ where: { id: event.childId } });
        return db.user.findMany({
          where: { familyId: child.familyId, role: "PARENT" },
        });

      case "MONTHLY_SUMMARY":
        return db.user.findMany({
          where: { familyId: event.familyId, role: "PARENT" },
        });
    }
  },

  private shouldEmail(event: NotificationEvent): boolean {
    // Only email for significant events, not every token award
    return ["GOAL_REACHED", "MONTHLY_SUMMARY"].includes(event.type);
  },
};""")

# ============================================================================
# SECTION 11: SCHEDULER SYSTEM
# ============================================================================
add_heading("11. Scheduler System", level=1)

add_para("Vercel Cron triggers scheduled jobs. Heavy processing runs on BullMQ + Redis.", bold=True, size=12)

add_code("""// vercel.json
{
  "crons": [
    { "path": "/api/cron/monthly-summary", "schedule": "0 8 1 * *" },
    { "path": "/api/cron/goal-reset", "schedule": "0 0 * * 0" },
    { "path": "/api/cron/streak-calc", "schedule": "0 1 * * *" },
    { "path": "/api/cron/allowance-payout", "schedule": "0 9 * * 1" }
  ]
}

// src/app/api/cron/monthly-summary/route.ts
import { NextRequest, NextResponse } from "next/server";
import { MonthlySummaryJob } from "@/server/jobs/monthly-summary.job";

export async function GET(req: NextRequest) {
  // Verify Vercel Cron signature
  const authHeader = req.headers.get("authorization");
  if (authHeader !== `Bearer ${process.env.CRON_SECRET}`) {
    return NextResponse.json({ error: "Unauthorized" }, { status: 401 });
  }

  try {
    await MonthlySummaryJob.run();
    return NextResponse.json({ ok: true });
  } catch (err) {
    console.error("Monthly summary job failed:", err);
    return NextResponse.json({ error: "Job failed" }, { status: 500 });
  }
}

// src/server/jobs/monthly-summary.job.ts
import { db } from "@/lib/db";
import { NotificationService } from "@/server/services/notification.service";
import { ReportService } from "@/server/services/report.service";

export const MonthlySummaryJob = {
  async run() {
    const families = await db.family.findMany({ where: { deletedAt: null } });

    for (const family of families) {
      try {
        const report = await ReportService.generateMonthlySummary(family.id);

        await NotificationService.notify({
          type: "MONTHLY_SUMMARY",
          familyId: family.id,
        });

        // Idempotency: mark as sent
        await db.familySettings.update({
          where: { familyId: family.id },
          data: { lastMonthlySummaryAt: new Date() },
        });
      } catch (err) {
        console.error(`Summary for family ${family.id} failed:`, err);
        // Continue to next family — one failure shouldn't block others
      }
    }
  },
};""")

# ============================================================================
# SECTION 12: REPORTING SYSTEM
# ============================================================================
page_break()
add_heading("12. Reporting System", level=1)

add_para("CSV/PDF exports + monthly reports via background jobs.", bold=True, size=12)

add_code("""// src/server/services/report.service.ts
import { db } from "@/lib/db";
import { formatCurrency } from "@/lib/currency";
import { emailSender } from "@/server/emails/sender";

export const ReportService = {
  async generateMonthlySummary(familyId: string) {
    const now = new Date();
    const monthStart = new Date(now.getFullYear(), now.getMonth() - 1, 1);
    const monthEnd = new Date(now.getFullYear(), now.getMonth(), 0, 23, 59, 59);

    const family = await db.family.findUnique({ where: { id: familyId } });
    const children = await db.child.findMany({ where: { familyId, deletedAt: null } });

    const summary = {
      familyName: family.name,
      month: monthStart.toLocaleDateString("en-US", { month: "long", year: "numeric" }),
      currency: family.currency,
      children: [] as any[],
      totalSaved: 0,
      totalSpent: 0,
      goalsReached: 0,
    };

    for (const child of children) {
      const saved = await db.transaction.aggregate({
        where: {
          childId: child.id,
          type: "SAVE",
          timestamp: { gte: monthStart, lte: monthEnd },
        },
        _sum: { amount: true },
      });

      const spent = await db.spendingEntry.aggregate({
        where: {
          ownerId: child.id,
          timestamp: { gte: monthStart, lte: monthEnd },
        },
        _sum: { amount: true },
      });

      const goalsReached = await db.goal.count({
        where: {
          ownerId: child.id,
          currentAmount: { gte: db.goal.fields.targetAmount },
        },
      });

      summary.children.push({
        name: child.name,
        saved: saved._sum.amount ?? 0,
        spent: spent._sum.amount ?? 0,
        goalsReached,
      });
      summary.totalSaved += saved._sum.amount ?? 0;
      summary.totalSpent += spent._sum.amount ?? 0;
      summary.goalsReached += goalsReached;
    }

    return summary;
  },

  async exportTransactionsCSV(familyId: string, opts: { childId?: string; from?: Date; to?: Date }) {
    const transactions = await db.transaction.findMany({
      where: {
        familyId,
        childId: opts.childId,
        timestamp: { gte: opts.from, lte: opts.to },
      },
      orderBy: { timestamp: "desc" },
    });

    const headers = ["Date", "Child", "Type", "Amount (UGX)", "Tokens", "Note"];
    const rows = transactions.map((t) => [
      new Date(t.timestamp).toISOString(),
      t.child.name,
      t.type,
      t.amount.toString(),
      t.tokenDelta.toString(),
      t.note,
    ]);

    return [headers, ...rows].map((r) => r.join(",")).join("\\n");
  },

  async emailMonthlyReport(familyId: string) {
    const summary = await this.generateMonthlySummary(familyId);
    const parents = await db.user.findMany({
      where: { familyId, role: "PARENT" },
    });

    for (const parent of parents) {
      await emailSender.send({
        to: parent.email,
        subject: `${summary.month} Summary for ${summary.familyName}`,
        template: "monthly-summary",
        data: summary,
      });
    }
  },
};""")

# ============================================================================
# SECTION 13: FILE UPLOAD SYSTEM
# ============================================================================
page_break()
add_heading("13. File Upload System", level=1)

add_para("Presigned S3 uploads + image processing via Cloudinary.", bold=True, size=12)

add_code("""// src/app/api/uploads/presign/route.ts
import { NextRequest, NextResponse } from "next/server";
import { S3Client, PutObjectCommand } from "@aws-sdk/client-s3";
import { getSignedUrl } from "@aws-sdk/s3-request-presigner";
import { requireAuth } from "@/server/permissions/guards";
import { z } from "zod";

const s3 = new S3Client({
  region: process.env.S3_REGION,
  credentials: {
    accessKeyId: process.env.S3_ACCESS_KEY!,
    secretAccessKey: process.env.S3_SECRET_KEY!,
  },
});

const presignSchema = z.object({
  filename: z.string(),
  mimeType: z.string().regex(/^image\\/(jpeg|png|webp)$/),
  size: z.number().int().positive().max(5_000_000),  // 5MB max
  purpose: z.enum(["avatar", "document"]),
});

export async function POST(req: NextRequest) {
  const user = await requireAuth();

  const body = await req.json();
  const input = presignSchema.parse(body);

  const key = `uploads/${user.familyId}/${input.purpose}/${crypto.randomUUID()}-${input.filename}`;

  const command = new PutObjectCommand({
    Bucket: process.env.S3_BUCKET!,
    Key: key,
    ContentType: input.mimeType,
    ContentLength: input.size,
  });

  const uploadUrl = await getSignedUrl(s3, command, { expiresIn: 300 });

  return NextResponse.json({ uploadUrl, key });
}

// src/app/api/uploads/confirm/route.ts
export async function POST(req: NextRequest) {
  const user = await requireAuth();

  const { key, purpose } = await req.json();

  // For avatars: generate thumbnail via Cloudinary
  let thumbnailUrl: string | undefined;
  if (purpose === "avatar") {
    thumbnailUrl = await generateThumbnail(key, 64);
  }

  // Record in database
  const upload = await db.fileUpload.create({
    data: {
      familyId: user.familyId,
      uploaderId: user.id,
      filename: key.split("/").pop()!,
      mimeType: "image/jpeg",
      size: 0,  // Updated via S3 event
      storageKey: key,
      storageUrl: `${process.env.CDN_BASE_URL}/${key}`,
      thumbnailUrl,
    },
  });

  return NextResponse.json({ upload });
}""")

# ============================================================================
# SECTION 14: AUDIT TRAIL SYSTEM
# ============================================================================
page_break()
add_heading("14. Audit Trail System", level=1)

add_para("Immutable audit log capturing who/what/when/where/before/after for every state change.", bold=True, size=12)

add_code("""// src/server/services/audit.service.ts
import { db } from "@/lib/db";
import { auth } from "@/lib/auth";
import { headers } from "next/headers";

export const AuditService = {
  async log(input: {
    userId?: string;
    action: string;
    entityType: string;
    entityId: string;
    before?: any;
    after?: any;
  }) {
    const session = await auth();
    const headerList = headers();

    await db.auditLog.create({
      data: {
        userId: input.userId ?? session?.user?.id,
        familyId: session?.user?.familyId,
        action: input.action,
        entityType: input.entityType,
        entityId: input.entityId,
        before: input.before ?? null,
        after: input.after ?? null,
        ipAddress: headerList.get("x-forwarded-for") ?? null,
        userAgent: headerList.get("user-agent") ?? null,
      },
    });
  },

  async query(filter: {
    userId?: string;
    familyId?: string;
    action?: string;
    entityType?: string;
    entityId?: string;
    from?: Date;
    to?: Date;
    limit?: number;
    offset?: number;
  }) {
    return db.auditLog.findMany({
      where: {
        userId: filter.userId,
        familyId: filter.familyId,
        action: filter.action as any,
        entityType: filter.entityType,
        entityId: filter.entityId,
        createdAt: { gte: filter.from, lte: filter.to },
      },
      orderBy: { createdAt: "desc" },
      take: filter.limit ?? 50,
      skip: filter.offset ?? 0,
    });
  },

  async exportCSV(filter: Parameters<typeof this.query>[0]): Promise<string> {
    const logs = await this.query({ ...filter, limit: 10000 });
    const headers = ["Timestamp", "User", "Action", "Entity", "Entity ID", "IP", "Details"];
    const rows = logs.map((l) => [
      l.createdAt.toISOString(),
      l.userId ?? "",
      l.action,
      l.entityType,
      l.entityId,
      l.ipAddress ?? "",
      JSON.stringify({ before: l.before, after: l.after }),
    ]);
    return [headers, ...rows].map((r) => r.map((c) => `"${c}"`).join(",")).join("\\n");
  },
};""")

# ============================================================================
# SECTION 15: SECURITY ARCHITECTURE
# ============================================================================
page_break()
add_heading("15. Security Architecture", level=1)

add_para("Security checklist + implementation:", bold=True, size=12)

add_table(
    ["Layer", "Threat", "Mitigation"],
    [
        ["Transport", "MITM", "HTTPS-only, HSTS header (max-age 1 year)"],
        ["Auth", "Brute force", "Rate limit 5 logins/15min per IP, lockout after 5 failures"],
        ["Auth", "Credential stuffing", "bcrypt cost 12, password breach check via HaveIBeenPwned API"],
        ["Session", "Session hijacking", "HttpOnly + Secure + SameSite=Lax cookies, rotate session token on login"],
        ["Session", "CSRF", "SameSite cookies + double-submit token for mutations"],
        ["Input", "SQL injection", "Prisma parameterized queries (never raw SQL)"],
        ["Input", "XSS", "React auto-escaping + DOMPurify for user HTML"],
        ["Input", "NoSQL injection", "Zod validation on every API + server action input"],
        ["Authz", "IDOR", "requireFamilyAccess + requireChildAccess guards on every route"],
        ["Authz", "Privilege escalation", "Role checks via requirePermission before every mutation"],
        ["Files", "Malicious upload", "MIME type validation + file size limit + virus scan (ClamAV)"],
        ["Rate", "API abuse", "Redis-backed rate limiter: 100 req/min per user, 1000/min per IP"],
        ["Secrets", "Leaked keys", "Environment validation via zod, no secrets in git"],
        ["Deps", "Vulnerable packages", "Dependabot + npm audit in CI"],
        ["Headers", "Clickjacking", "X-Frame-Options: DENY"],
        ["Headers", "MIME sniffing", "X-Content-Type-Options: nosniff"],
    ],
)

add_code("""// src/lib/rate-limit.ts
import { Redis } from "@upstash/redis";

const redis = new Redis({
  url: process.env.UPSTASH_REDIS_URL!,
  token: process.env.UPSTASH_REDIS_TOKEN!,
});

export async function rateLimit(
  key: string,
  limit: number,
  windowSeconds: number,
): Promise<{ success: boolean; remaining: number }> {
  const count = await redis.incr(key);
  if (count === 1) {
    await redis.expire(key, windowSeconds);
  }
  return {
    success: count <= limit,
    remaining: Math.max(0, limit - count),
  };
}

// src/lib/env.ts
import { z } from "zod";

const envSchema = z.object({
  DATABASE_URL: z.string().url(),
  NEXTAUTH_SECRET: z.string().min(32),
  NEXTAUTH_URL: z.string().url(),
  SMTP_HOST: z.string(),
  SMTP_USER: z.string(),
  SMTP_PASS: z.string(),
  EMAIL_FROM: z.string().email(),
  S3_REGION: z.string(),
  S3_ACCESS_KEY: z.string(),
  S3_SECRET_KEY: z.string(),
  S3_BUCKET: z.string(),
  CDN_BASE_URL: z.string().url(),
  UPSTASH_REDIS_URL: z.string().url(),
  UPSTASH_REDIS_TOKEN: z.string(),
  CRON_SECRET: z.string().min(32),
  SENTRY_DSN: z.string().optional(),
});

export const env = envSchema.parse(process.env);""")

# ============================================================================
# SECTION 16: DEPLOYMENT ARCHITECTURE
# ============================================================================
page_break()
add_heading("16. Deployment Architecture", level=1)

add_para("Vercel + PostgreSQL (Neon/Supabase) + S3 + Upstash Redis + Resend Email.", bold=True, size=12)

add_code("""# .env.example

# Database
DATABASE_URL="postgresql://user:pass@host:5432/planned?schema=public"

# Auth
NEXTAUTH_SECRET="generate-with-openssl-rand-base64-32"
NEXTAUTH_URL="https://planned.app"

# Email (Resend)
RESEND_API_KEY="re_xxxxx"
EMAIL_FROM="Planned <hello@planned.app>"

# File Storage (S3 / Cloudinary)
S3_REGION="us-east-1"
S3_ACCESS_KEY="AKIAXXXXX"
S3_SECRET_KEY="xxxxx"
S3_BUCKET="planned-uploads"
CDN_BASE_URL="https://cdn.planned.app"

# Cache + Queue (Upstash Redis)
UPSTASH_REDIS_URL="https://xxx.upstash.io"
UPSTASH_REDIS_TOKEN="xxxxx"

# Cron
CRON_SECRET="generate-with-openssl-rand-base64-32"

# Observability
SENTRY_DSN="https://xxx@sentry.io/xxx"

# Feature flags
NEXT_PUBLIC_APP_URL="https://planned.app"
NEXT_PUBLIC_STRIPE_PUBLISHABLE_KEY="pk_test_xxx"  # Future monetization""")

add_para("CI/CD Pipeline (GitHub Actions):", bold=True, size=12)
add_code("""# .github/workflows/deploy.yml
name: Deploy

on:
  push:
    branches: [main]
  pull_request:
    branches: [main]

jobs:
  test:
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v4
      - uses: oven-sh/setup-bun@v1
      - run: bun install
      - run: bun run lint
      - run: bun run typecheck
      - run: bun run test
      - run: bun run prisma generate
      - run: bun run prisma migrate deploy
        env:
          DATABASE_URL: \${{ secrets.DATABASE_URL }}

  deploy:
    needs: test
    if: github.ref == 'refs/heads/main'
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v4
      - uses: amondnet/vercel-action@v25
        with:
          vercel-token: \${{ secrets.VERCEL_TOKEN }}
          vercel-org-id: \${{ secrets.VERCEL_ORG_ID }}
          vercel-project-id: \${{ secrets.VERCEL_PROJECT_ID }}
          vercel-args: '--prod'""")

# ============================================================================
# SECTION 17: ENVIRONMENT VARIABLES
# ============================================================================
page_break()
add_heading("17. Environment Variables", level=1)

add_table(
    ["Variable", "Purpose", "Required"],
    [
        ["DATABASE_URL", "PostgreSQL connection string", "Yes"],
        ["NEXTAUTH_SECRET", "JWT + session encryption (32+ chars)", "Yes"],
        ["NEXTAUTH_URL", "App URL for callbacks", "Yes"],
        ["RESEND_API_KEY", "Email sending", "Yes"],
        ["EMAIL_FROM", "From address", "Yes"],
        ["S3_REGION / S3_ACCESS_KEY / S3_SECRET_KEY / S3_BUCKET", "File storage", "Yes"],
        ["CDN_BASE_URL", "CloudFront/Cloudinary URL", "Yes"],
        ["UPSTASH_REDIS_URL / UPSTASH_REDIS_TOKEN", "Cache + rate limit + queue", "Yes"],
        ["CRON_SECRET", "Protects cron endpoints", "Yes"],
        ["SENTRY_DSN", "Error tracking", "Optional"],
        ["NEXT_PUBLIC_STRIPE_PUBLISHABLE_KEY", "Future payments", "Optional"],
    ],
)

# ============================================================================
# SECTION 18: MIGRATION PLAN
# ============================================================================
add_heading("18. Migration Plan", level=1)

add_para("Phase 1: SQLite → PostgreSQL (1 day)", bold=True, size=12)
add_bullet("Update DATABASE_URL to PostgreSQL connection string")
add_bullet("Run prisma migrate dev to create migration from schema")
add_bullet("Write one-time data migration script: read from SQLite, insert into PostgreSQL")
add_bullet("Run prisma migrate deploy on production")

add_para("Phase 2: Add authentication (3 days)", bold=True, size=12)
add_bullet("Configure NextAuth with Credentials + Email providers")
add_bullet("Add middleware for route protection")
add_bullet("Migrate existing Family data to have a User (parent) account")
add_bullet("Add login/register/forgot-password pages")

add_para("Phase 3: Multi-tenant isolation (2 days)", bold=True, size=12)
add_bullet("Add familyId to all existing entities (they already have it)")
add_bullet("Add requireFamilyAccess guard to every route + server action")
add_bullet("Verify no cross-family data leakage with integration tests")

add_para("Phase 4: File uploads (2 days)", bold=True, size=12)
add_bullet("Set up S3 bucket + CloudFront distribution")
add_bullet("Build presigned upload API")
add_bullet("Migrate avatar photos from base64 (in DB) to S3 URLs")

add_para("Phase 5: Notifications + Scheduler (3 days)", bold=True, size=12)
add_bullet("Set up Resend for email")
add_bullet("Build notification service + in-app dropdown")
add_bullet("Configure Vercel Cron for monthly summaries + goal resets")

add_para("Phase 6: Audit + Reports (2 days)", bold=True, size=12)
add_bullet("Add audit log to every mutation")
add_bullet("Build CSV export endpoint")
add_bullet("Build monthly PDF report generator")

# ============================================================================
# SECTION 19: SCALING STRATEGY
# ============================================================================
page_break()
add_heading("19. Scaling Strategy", level=1)

add_para("Target: 10,000 families, 50,000 children, millions of transactions.", bold=True, size=12)

add_table(
    ["Bottleneck", "Solution"],
    [
        ["Database reads", "Read replicas + React Query caching (30s staleTime)"],
        ["Database writes", "Connection pooling via PgBouncer, optimistic concurrency"],
        ["Large tables (Transaction)", "Partition by year + familyId. Archive rows > 2 years to cold storage."],
        ["File storage", "S3 + CloudFront CDN. No files in database."],
        ["Background jobs", "BullMQ + Redis. Vercel Cron triggers, workers process."],
        ["Search", "PostgreSQL full-text search initially. Migrate to Meilisearch if needed."],
        ["Rate limiting", "Redis-backed sliding window. 100 req/min/user."],
        ["Memory", "Vercel serverless functions: 1GB max. Stream large responses."],
    ],
)

add_para("Partitioning strategy for Transaction table (50M+ rows):", bold=True, size=12)
add_code("""-- PostgreSQL native partitioning by year + familyId hash
CREATE TABLE transaction (
  id UUID,
  family_id UUID,
  child_id UUID,
  type TEXT,
  amount INTEGER,
  timestamp TIMESTAMPTZ,
  ...
) PARTITION BY RANGE (timestamp);

CREATE TABLE transaction_2025 PARTITION OF transaction
  FOR VALUES FROM ('2025-01-01') TO ('2026-01-01');

CREATE TABLE transaction_2026 PARTITION OF transaction
  FOR VALUES FROM ('2026-01-01') TO ('2027-01-01');

-- Sub-partition by familyId hash for write scaling
CREATE TABLE transaction_2025_f0 PARTITION OF transaction_2025
  FOR VALUES WITH (modulus 4, remainder 0);
-- ... etc""")

# ============================================================================
# SECTION 20: IMPLEMENTATION ROADMAP
# ============================================================================
add_heading("20. Implementation Roadmap", level=1)

add_para("12-week roadmap to production:", bold=True, size=12)

add_table(
    ["Week", "Focus", "Deliverables"],
    [
        ["1-2", "Database + Auth", "PostgreSQL migration, NextAuth, login/register, RBAC guards"],
        ["3", "Multi-tenant isolation", "familyId on all entities, guards on all routes"],
        ["4", "Repository + Service layers", "Refactor all DB access behind repositories, extract services"],
        ["5", "Token economy (black-box)", "TokenService interface, wire up award/redeem via interface"],
        ["6", "File uploads", "S3 presigned uploads, avatar migration from base64"],
        ["7", "Notifications", "In-app + email, notification preferences, Resend integration"],
        ["8", "Scheduler + Jobs", "Vercel Cron, monthly summaries, goal resets, streak calc"],
        ["9", "Reporting", "CSV export, PDF monthly report, email delivery"],
        ["10", "Audit trail", "AuditService on every mutation, audit log UI, CSV export"],
        ["11", "Security hardening", "Rate limiting, security headers, pentest, dependency audit"],
        ["12", "Observability + Launch", "Sentry, structured logging, health checks, production deploy"],
    ],
)

add_para("Success Criteria:", bold=True, size=12)
add_bullet("10,000 families can use the app simultaneously without degradation")
add_bullet("Every mutation is audited + persists to PostgreSQL")
add_bullet("Every route is protected by auth + RBAC")
add_bullet("Every file upload goes through presigned S3 URLs (not base64 in DB)")
add_bullet("Every notification can be delivered via in-app + email + push")
add_bullet("Every report can be generated + emailed as PDF/CSV")
add_bullet("Every background job is idempotent + retry-safe")
add_bullet("Zero P0 security vulnerabilities (verified by pentest)")

# Save final
doc.save(str(OUT))
print(f"Final document saved: {OUT}")
print(f"Total size: {OUT.stat().st_size:,} bytes")
