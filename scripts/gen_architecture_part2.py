"""Generate Planned Production Architecture DOCX — Part 2: Sections 4-9 (Auth, Authorization, API, State, Services, Repository)."""
import sys, os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path("/home/z/my-project/public/Planned_Production_Architecture.docx")
doc = Document(str(OUT))  # Load existing

GOLD = RGBColor(201, 168, 76)
DARK = RGBColor(30, 30, 30)
GRAY = RGBColor(120, 120, 120)
CODE_COLOR = RGBColor(40, 80, 40)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GOLD if level == 1 else DARK
    return h

def add_para(text, bold=False, italic=False, color=None, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    r.font.size = Pt(size)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_code(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.line_spacing = 1.15
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F0")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "CCCCCC")
        pBdr.append(border)
    pPr.append(pBdr)
    r = p.add_run(code_text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = CODE_COLOR
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def page_break():
    doc.add_page_break()

# ============================================================================
# SECTION 4: AUTHENTICATION ARCHITECTURE
# ============================================================================
page_break()
add_heading("4. Authentication Architecture", level=1)

add_para("Auth.js (NextAuth v5) configuration with Credentials + Email providers, session management, and brute-force protection:", bold=True, size=12)

add_code("""// src/lib/auth.ts
import NextAuth from "next-auth";
import Credentials from "next-auth/providers/credentials";
import Email from "next-auth/providers/email";
import { PrismaAdapter } from "@auth/prisma-adapter";
import bcrypt from "bcryptjs";
import { db } from "@/lib/db";
import { rateLimit } from "@/lib/rate-limit";
import { z } from "zod";

const credentialsSchema = z.object({
  email: z.string().email(),
  password: z.string().min(8),
  rememberMe: z.boolean().optional(),
});

export const { handlers, auth, signIn, signOut } = NextAuth({
  adapter: PrismaAdapter(db),
  session: {
    strategy: "database",  // DB-backed sessions for instant invalidation
    maxAge: 30 * 24 * 60 * 60,  // 30 days
    updateAge: 24 * 60 * 60,    // Update session every 24h
  },
  pages: {
    signIn: "/login",
    signOut: "/logout",
    error: "/login",
    verifyRequest: "/verify-request",
    newUser: "/register",
  },
  providers: [
    Credentials({
      name: "credentials",
      credentials: {
        email: { label: "Email", type: "email" },
        password: { label: "Password", type: "password" },
      },
      async authorize(credentials, request) {
        const parsed = credentialsSchema.safeParse(credentials);
        if (!parsed.success) return null;

        const { email, password } = parsed.data;

        // Rate limit: 5 attempts per IP per 15 minutes
        const ip = request.headers.get("x-forwarded-for") ?? "unknown";
        const { success } = await rateLimit(`login:${ip}`, 5, 900);
        if (!success) {
          throw new Error("Too many login attempts. Try again in 15 minutes.");
        }

        const user = await db.user.findUnique({
          where: { email: email.toLowerCase() },
        });

        if (!user || !user.passwordHash) return null;

        // Check account lockout
        if (user.lockedUntil && user.lockedUntil > new Date()) {
          throw new Error("Account locked. Try again later.");
        }

        // Verify password
        const valid = await bcrypt.compare(password, user.passwordHash);
        if (!valid) {
          // Increment failed attempts, lock after 5
          await db.user.update({
            where: { id: user.id },
            data: {
              failedLoginAttempts: { increment: 1 },
              lockedUntil:
                user.failedLoginAttempts + 1 >= 5
                  ? new Date(Date.now() + 15 * 60 * 1000)
                  : null,
            },
          });
          return null;
        }

        // Reset failed attempts on success
        await db.user.update({
          where: { id: user.id },
          data: {
            failedLoginAttempts: 0,
            lockedUntil: null,
          },
        });

        return {
          id: user.id,
          email: user.email,
          name: user.name,
          role: user.role,
          familyId: user.familyId,
        };
      },
    }),
    Email({
      server: {
        host: process.env.SMTP_HOST,
        port: 587,
        auth: { user: process.env.SMTP_USER, pass: process.env.SMTP_PASS },
      },
      from: process.env.EMAIL_FROM,
    }),
  ],
  callbacks: {
    async jwt({ token, user }) {
      if (user) {
        token.role = user.role;
        token.familyId = user.familyId;
      }
      return token;
    },
    async session({ session, token }) {
      if (session.user) {
        session.user.id = token.sub!;
        session.user.role = token.role as Role;
        session.user.familyId = token.familyId as string;
      }
      return session;
    },
    authorized({ auth, request }) {
      // Protect all /dashboard and /api routes except /api/auth
      const pathname = request.nextUrl.pathname;
      if (pathname.startsWith("/dashboard") && !auth) return false;
      if (pathname.startsWith("/api/") && !pathname.startsWith("/api/auth") && !auth) {
        return false;
      }
      return true;
    },
  },
  events: {
    async signIn({ user, isNewUser }) {
      // Audit log
      await db.auditLog.create({
        data: {
          userId: user.id,
          action: "LOGIN",
          entityType: "User",
          entityId: user.id,
          ipAddress: "",
        },
      });
    },
  },
});""")

add_para("Middleware (route protection + CSRF + rate limiting):", bold=True, size=12)
add_code("""// src/middleware.ts
import { auth } from "@/lib/auth";
import { NextResponse } from "next/server";

export default auth((req) => {
  const { pathname } = req.nextUrl;
  const isLoggedIn = !!req.auth;

  // Public routes
  const publicRoutes = ["/", "/login", "/register", "/forgot-password",
                        "/api/auth", "/api/health"];
  if (publicRoutes.some(r => pathname.startsWith(r))) {
    return NextResponse.next();
  }

  // Protected routes — require auth
  if (!isLoggedIn) {
    return NextResponse.redirect(new URL("/login", req.url));
  }

  // Role-based route protection
  const role = req.auth.user.role;
  if (pathname.startsWith("/dashboard") && role === "CHILD") {
    return NextResponse.redirect(new URL("/child", req.url));
  }
  if (pathname.startsWith("/child") && role === "PARENT") {
    return NextResponse.redirect(new URL("/dashboard", req.url));
  }

  // Security headers
  const response = NextResponse.next();
  response.headers.set("X-Frame-Options", "DENY");
  response.headers.set("X-Content-Type-Options", "nosniff");
  response.headers.set("Referrer-Policy", "strict-origin-when-cross-origin");
  response.headers.set("Permissions-Policy", "camera=(), microphone=(), geolocation=()");
  response.headers.set("Strict-Transport-Security", "max-age=31536000; includeSubDomains");

  return response;
});

export const config = {
  matcher: ["/((?!_next/static|_next/image|favicon.ico).*)"],
};""")

add_para("Password Reset Flow:", bold=True, size=12)
add_code("""// src/features/auth/actions.ts
"use server";

import { z } from "zod";
import bcrypt from "bcryptjs";
import { db } from "@/lib/db";
import { sendPasswordResetEmail } from "@/server/emails/sender";
import { generateToken } from "@/lib/crypto";

const forgotPasswordSchema = z.object({
  email: z.string().email(),
});

export async function forgotPassword(formData: FormData) {
  const { email } = forgotPasswordSchema.parse({
    email: formData.get("email"),
  });

  const user = await db.user.findUnique({ where: { email } });
  if (!user) {
    // Don't reveal whether email exists — security best practice
    return { success: true };
  }

  const token = generateToken();
  const expires = new Date(Date.now() + 60 * 60 * 1000); // 1 hour

  await db.passwordResetToken.create({
    data: { userId: user.id, token, expires },
  });

  await sendPasswordResetEmail(email, token);

  return { success: true };
}

const resetPasswordSchema = z.object({
  token: z.string(),
  password: z.string().min(8).max(100),
});

export async function resetPassword(formData: FormData) {
  const { token, password } = resetPasswordSchema.parse({
    token: formData.get("token"),
    password: formData.get("password"),
  });

  const resetToken = await db.passwordResetToken.findUnique({
    where: { token },
  });

  if (!resetToken || resetToken.expires < new Date()) {
    return { error: "Invalid or expired token" };
  }

  const passwordHash = await bcrypt.hash(password, 12);

  await db.$transaction([
    db.user.update({
      where: { id: resetToken.userId },
      data: { passwordHash },
    }),
    db.passwordResetToken.delete({ where: { id: resetToken.id } }),
    // Invalidate all existing sessions
    db.session.deleteMany({ where: { userId: resetToken.userId } }),
  ]);

  return { success: true };
}""")

add_para("RBAC Summary:", bold=True, size=12)
add_table(
    ["Role", "Routes", "Permissions"],
    [
        ["ADMIN", "/admin/*", "Manage all families, view system metrics, manage users"],
        ["PARENT", "/dashboard/*", "Manage family, children, goals, investments, tokens, settings"],
        ["CHILD", "/child/*", "View own dashboard, log spending, redeem tokens, view own goals"],
    ],
)

page_break()

# ============================================================================
# SECTION 5: AUTHORIZATION ARCHITECTURE
# ============================================================================
add_heading("5. Authorization Architecture", level=1)

add_para("Permission matrix + server-side guards:", bold=True, size=12)

add_code("""// src/server/permissions/roles.ts
export const ROLES = {
  ADMIN: "ADMIN",
  PARENT: "PARENT",
  CHILD: "CHILD",
} as const;

export type Role = typeof ROLES[keyof typeof ROLES];

// src/server/permissions/permissions.ts
export const PERMISSIONS = {
  // Family
  FAMILY_VIEW: "family:view",
  FAMILY_MANAGE: "family:manage",

  // Children
  CHILD_VIEW: "child:view",
  CHILD_CREATE: "child:create",
  CHILD_UPDATE: "child:update",
  CHILD_DELETE: "child:delete",

  // Transactions
  TRANSACTION_VIEW: "transaction:view",
  TRANSACTION_CREATE: "transaction:create",
  TRANSACTION_DELETE: "transaction:delete",

  // Goals
  GOAL_VIEW: "goal:view",
  GOAL_CREATE: "goal:create",
  GOAL_UPDATE: "goal:update",
  GOAL_DELETE: "goal:delete",
  GOAL_CONTRIBUTE: "goal:contribute",

  // Investments
  INVESTMENT_VIEW: "investment:view",
  INVESTMENT_MANAGE: "investment:manage",

  // Tokens
  TOKEN_AWARD: "token:award",
  TOKEN_REDEEM: "token:redeem",
  TOKEN_VIEW_LEDGER: "token:view:ledger",

  // Reports
  REPORT_VIEW: "report:view",
  REPORT_GENERATE: "report:generate",

  // Settings
  SETTINGS_VIEW: "settings:view",
  SETTINGS_MANAGE: "settings:manage",

  // Uploads
  UPLOAD_AVATAR: "upload:avatar",
} as const;

export type Permission = typeof PERMISSIONS[keyof typeof PERMISSIONS];

// Role → Permission matrix
export const ROLE_PERMISSIONS: Record<Role, Permission[]> = {
  ADMIN: Object.values(PERMISSIONS), // All permissions

  PARENT: [
    PERMISSIONS.FAMILY_VIEW,
    PERMISSIONS.FAMILY_MANAGE,
    PERMISSIONS.CHILD_VIEW,
    PERMISSIONS.CHILD_CREATE,
    PERMISSIONS.CHILD_UPDATE,
    PERMISSIONS.CHILD_DELETE,
    PERMISSIONS.TRANSACTION_VIEW,
    PERMISSIONS.TRANSACTION_CREATE,
    PERMISSIONS.TRANSACTION_DELETE,
    PERMISSIONS.GOAL_VIEW,
    PERMISSIONS.GOAL_CREATE,
    PERMISSIONS.GOAL_UPDATE,
    PERMISSIONS.GOAL_DELETE,
    PERMISSIONS.GOAL_CONTRIBUTE,
    PERMISSIONS.INVESTMENT_VIEW,
    PERMISSIONS.INVESTMENT_MANAGE,
    PERMISSIONS.TOKEN_AWARD,
    PERMISSIONS.TOKEN_VIEW_LEDGER,
    PERMISSIONS.REPORT_VIEW,
    PERMISSIONS.REPORT_GENERATE,
    PERMISSIONS.SETTINGS_VIEW,
    PERMISSIONS.SETTINGS_MANAGE,
    PERMISSIONS.UPLOAD_AVATAR,
  ],

  CHILD: [
    PERMISSIONS.CHILD_VIEW,        // Only own data
    PERMISSIONS.TRANSACTION_VIEW,  // Only own transactions
    PERMISSIONS.TRANSACTION_CREATE,// Only save/redeem for self
    PERMISSIONS.GOAL_VIEW,         // Only revealed + own private
    PERMISSIONS.GOAL_CONTRIBUTE,   // Only to own goals
    PERMISSIONS.INVESTMENT_VIEW,   // Only own investments
    PERMISSIONS.TOKEN_REDEEM,      // Only own tokens
    PERMISSIONS.REPORT_VIEW,       // Only own reports
    PERMISSIONS.UPLOAD_AVATAR,     // Only own avatar
  ],
};""")

add_code("""// src/server/permissions/guards.ts
import { PERMISSIONS, ROLE_PERMISSIONS, type Permission, type Role } from "./permissions";
import { auth } from "@/lib/auth";
import { db } from "@/lib/db";
import { AuthorizationError } from "@/lib/errors";

export async function getCurrentUser() {
  const session = await auth();
  if (!session?.user) return null;
  return session.user;
}

export async function requireAuth() {
  const user = await getCurrentUser();
  if (!user) {
    throw new AuthorizationError("Authentication required");
  }
  return user;
}

export async function requirePermission(permission: Permission) {
  const user = await requireAuth();
  const allowed = ROLE_PERMISSIONS[user.role as Role];
  if (!allowed.includes(permission)) {
    throw new AuthorizationError(`Missing permission: ${permission}`);
  }
  return user;
}

// Ensure the user can only access resources within their own family.
export async function requireFamilyAccess(familyId: string) {
  const user = await requireAuth();
  if (user.role === "ADMIN") return user;
  if (user.familyId !== familyId) {
    throw new AuthorizationError("Cross-family access denied");
  }
  return user;
}

// Ensure a child can only access their own data.
export async function requireChildAccess(childId: string) {
  const user = await requireAuth();

  if (user.role === "ADMIN") return user;

  if (user.role === "PARENT") {
    // Parent can access any child in their family
    const child = await db.child.findUnique({
      where: { id: childId },
      select: { familyId: true },
    });
    if (!child || child.familyId !== user.familyId) {
      throw new AuthorizationError("Child not found in your family");
    }
    return user;
  }

  if (user.role === "CHILD") {
    // Child can only access their own record
    // (childId is stored on the User record for child users)
    if (user.childId !== childId) {
      throw new AuthorizationError("Cannot access another child's data");
    }
    return user;
  }

  throw new AuthorizationError("Invalid role");
}

// Ensure a child can only see revealed goals + their own private goals.
export async function filterGoalsByViewer<T extends { ownerId: string; visibility: string }>(
  goals: T[],
  viewerId: string,
  viewerRole: Role,
): Promise<T[]> {
  if (viewerRole === "ADMIN" || viewerRole === "PARENT") return goals;
  // CHILD: only revealed goals + their own private goals
  return goals.filter(
    (g) => g.visibility === "REVEALED" || g.ownerId === viewerId,
  );
}""")

add_para("Usage in Server Actions:", bold=True, size=12)
add_code("""// src/features/goals/actions.ts
"use server";

import { z } from "zod";
import { requirePermission, requireFamilyAccess } from "@/server/permissions/guards";
import { PERMISSIONS } from "@/server/permissions/permissions";
import { GoalService } from "@/server/services/goal.service";

const createGoalSchema = z.object({
  familyId: z.string().uuid(),
  ownerId: z.string().uuid(),
  ownerKind: z.enum(["parent", "child"]),
  title: z.string().min(1).max(80),
  type: z.enum(["save", "spend_less"]),
  cadence: z.enum(["weekly", "monthly", "annual"]),
  visibility: z.enum(["private", "revealed"]),
  targetAmount: z.number().int().positive(),
  note: z.string().max(120).optional(),
});

export async function createGoal(formData: FormData) {
  // 1. Authorize
  await requirePermission(PERMISSIONS.GOAL_CREATE);

  // 2. Validate
  const input = createGoalSchema.parse({
    familyId: formData.get("familyId"),
    ownerId: formData.get("ownerId"),
    // ... etc
  });

  // 3. Ensure family access
  await requireFamilyAccess(input.familyId);

  // 4. Delegate to service
  return GoalService.create(input);
}""")

page_break()

# ============================================================================
# SECTION 6: API ARCHITECTURE
# ============================================================================
add_heading("6. API Architecture", level=1)

add_para("REST API endpoints with Zod validation + typed responses:", bold=True, size=12)

add_para("Endpoint inventory:", bold=True, size=12)
add_table(
    ["Endpoint", "Method", "Purpose", "Auth Required"],
    [
        ["/api/auth/*", "GET/POST", "NextAuth handlers", "Public"],
        ["/api/families", "GET", "List user's family", "PARENT"],
        ["/api/families", "PATCH", "Update family settings", "PARENT"],
        ["/api/children", "GET", "List children in family", "PARENT"],
        ["/api/children", "POST", "Create child", "PARENT"],
        ["/api/children/:id", "GET", "Get child detail", "PARENT or CHILD(self)"],
        ["/api/children/:id", "PATCH", "Update child", "PARENT"],
        ["/api/children/:id", "DELETE", "Soft-delete child", "PARENT"],
        ["/api/accounts", "GET/POST", "List/create accounts", "PARENT"],
        ["/api/transactions", "GET", "List transactions (filterable)", "PARENT or CHILD(self)"],
        ["/api/transactions", "POST", "Create transaction (save/withdraw/invest/redeem)", "PARENT or CHILD(self)"],
        ["/api/goals", "GET/POST", "List/create goals", "PARENT or CHILD"],
        ["/api/goals/:id", "GET/PATCH/DELETE", "Manage single goal", "PARENT or CHILD(owner)"],
        ["/api/goals/:id/contribute", "POST", "Contribute to goal", "PARENT or CHILD"],
        ["/api/investments", "GET", "List investments", "PARENT or CHILD(self)"],
        ["/api/tokens/award", "POST", "Award tokens (parent only)", "PARENT"],
        ["/api/tokens/redeem", "POST", "Redeem tokens", "PARENT or CHILD(self)"],
        ["/api/uploads/presign", "POST", "Get presigned upload URL", "PARENT or CHILD"],
        ["/api/uploads/confirm", "POST", "Confirm upload completed", "PARENT or CHILD"],
        ["/api/reports/monthly", "GET", "Generate monthly report", "PARENT"],
        ["/api/reports/export", "GET", "Export CSV/Excel", "PARENT"],
        ["/api/notifications", "GET", "List notifications", "Any auth user"],
        ["/api/notifications/:id/read", "POST", "Mark as read", "Any auth user"],
    ],
)

add_para("Example: typed route handler with validation", bold=True, size=12)
add_code("""// src/app/api/transactions/route.ts
import { NextRequest, NextResponse } from "next/server";
import { z } from "zod";
import { requirePermission, requireFamilyAccess, requireChildAccess } from "@/server/permissions/guards";
import { PERMISSIONS } from "@/server/permissions/permissions";
import { TransactionService } from "@/server/services/transaction.service";
import { logger } from "@/lib/logger";

const createTransactionSchema = z.object({
  familyId: z.string().uuid(),
  childId: z.string().uuid(),
  type: z.enum(["save", "withdraw", "invest", "redeem", "parent_give"]),
  amount: z.number().int().positive().max(1_000_000_000),  // Max 1B UGX
  tokenDelta: z.number().int().min(0).default(0),
  accountId: z.string().uuid().optional(),
  investmentId: z.string().uuid().optional(),
  goalId: z.string().uuid().optional(),
  note: z.string().min(1).max(200),
});

export async function GET(req: NextRequest) {
  try {
    const user = await requirePermission(PERMISSIONS.TRANSACTION_VIEW);

    const { searchParams } = new URL(req.url);
    const familyId = searchParams.get("familyId");
    const childId = searchParams.get("childId");
    const type = searchParams.get("type");
    const limit = Math.min(Number(searchParams.get("limit") ?? "50"), 100);
    const offset = Number(searchParams.get("offset") ?? "0");

    if (familyId) await requireFamilyAccess(familyId);
    if (childId) await requireChildAccess(childId);

    const transactions = await TransactionService.list({
      familyId: familyId ?? user.familyId,
      childId: childId ?? undefined,
      type: type as any,
      limit,
      offset,
    });

    return NextResponse.json({ transactions, limit, offset });
  } catch (err) {
    logger.error({ err, path: req.url }, "GET /api/transactions failed");
    return NextResponse.json(
      { error: err instanceof Error ? err.message : "Internal error" },
      { status: err instanceof AuthorizationError ? 403 : 500 },
    );
  }
}

export async function POST(req: NextRequest) {
  try {
    const user = await requirePermission(PERMISSIONS.TRANSACTION_CREATE);

    const body = await req.json();
    const input = createTransactionSchema.parse(body);

    await requireFamilyAccess(input.familyId);
    await requireChildAccess(input.childId);

    const transaction = await TransactionService.create(input, user.id);

    return NextResponse.json({ transaction }, { status: 201 });
  } catch (err) {
    if (err instanceof z.ZodError) {
      return NextResponse.json(
        { error: "Validation failed", details: err.flatten() },
        { status: 400 },
      );
    }
    logger.error({ err }, "POST /api/transactions failed");
    return NextResponse.json(
      { error: err instanceof Error ? err.message : "Internal error" },
      { status: 500 },
    );
  }
}""")

page_break()

# ============================================================================
# SECTION 7: STATE MANAGEMENT ARCHITECTURE
# ============================================================================
add_heading("7. State Management Architecture", level=1)

add_para("Zustand is refactored to hold UI state ONLY. All domain data lives in React Query (server state cache).", bold=True, size=12)

add_para("What belongs in Zustand (UI state):", bold=True, size=12)
add_bullet("theme (onyx, ivory, blush, red)")
add_bullet("sidebar open/closed")
add_bullet("active modal (saveMoney, giveTokens, etc.)")
add_bullet("filters (selected child, date range, spending category)")
add_bullet("optimistic update markers (pending mutations)")

add_para("What does NOT belong in Zustand:", bold=True, size=12)
add_bullet("children array, transactions array, goals array — use React Query")
add_bullet("currentAmount balances — server-derived, cached by React Query")
add_bullet("anything that comes from the database")

add_code("""// src/stores/ui.store.ts
"use client";

import { create } from "zustand";
import { persist } from "zustand/middleware";

type Theme = "onyx" | "ivory" | "blush" | "red";
type ModalName = "saveMoney" | "giveTokens" | "addSpending" | "redeemTokens" | "investmentDetail" | "createGoal" | null;

interface UIState {
  // Theme
  theme: Theme;
  setTheme: (t: Theme) => void;

  // Sidebar
  sidebarOpen: boolean;
  toggleSidebar: () => void;
  setSidebarOpen: (open: boolean) => void;

  // Modals
  activeModal: ModalName;
  modalProps: Record<string, any>;
  openModal: (name: ModalName, props?: Record<string, any>) => void;
  closeModal: () => void;

  // Filters
  selectedChildId: string | null;
  setSelectedChild: (id: string | null) => void;
  dateRange: { start: Date; end: Date } | null;
  setDateRange: (range: { start: Date; end: Date } | null) => void;

  // Optimistic updates
  pendingMutations: Set<string>;
  addPendingMutation: (id: string) => void;
  removePendingMutation: (id: string) => void;
}

export const useUIStore = create<UIState>()(
  persist(
    (set, get) => ({
      theme: "onyx",
      setTheme: (theme) => set({ theme }),

      sidebarOpen: false,
      toggleSidebar: () => set((s) => ({ sidebarOpen: !s.sidebarOpen })),
      setSidebarOpen: (sidebarOpen) => set({ sidebarOpen }),

      activeModal: null,
      modalProps: {},
      openModal: (activeModal, modalProps = {}) => set({ activeModal, modalProps }),
      closeModal: () => set({ activeModal: null, modalProps: {} }),

      selectedChildId: null,
      setSelectedChild: (selectedChildId) => set({ selectedChildId }),
      dateRange: null,
      setDateRange: (dateRange) => set({ dateRange }),

      pendingMutations: new Set(),
      addPendingMutation: (id) =>
        set((s) => {
          const next = new Set(s.pendingMutations);
          next.add(id);
          return { pendingMutations: next };
        }),
      removePendingMutation: (id) =>
        set((s) => {
          const next = new Set(s.pendingMutations);
          next.delete(id);
          return { pendingMutations: next };
        }),
    }),
    {
      name: "planned-ui",
      partialize: (state) => ({ theme: state.theme }),  // Only persist theme
    },
  ),
);""")

add_para("Server state with React Query:", bold=True, size=12)
add_code("""// src/hooks/use-transactions.ts
"use client";

import { useQuery, useMutation, useQueryClient } from "@tanstack/react-query";
import { useUIStore } from "@/stores/ui.store";

export function useTransactions(childId?: string) {
  const selectedChildId = useUIStore((s) => s.selectedChildId);
  const effectiveChildId = childId ?? selectedChildId;

  return useQuery({
    queryKey: ["transactions", effectiveChildId],
    queryFn: async () => {
      const params = new URLSearchParams();
      if (effectiveChildId) params.set("childId", effectiveChildId);
      const res = await fetch(`/api/transactions?${params}`);
      if (!res.ok) throw new Error("Failed to load transactions");
      return res.json();
    },
    staleTime: 30_000,  // 30 seconds
  });
}

export function useCreateTransaction() {
  const queryClient = useQueryClient();
  const addPending = useUIStore((s) => s.addPendingMutation);
  const removePending = useUIStore((s) => s.removePendingMutation);

  return useMutation({
    mutationFn: async (input: CreateTransactionInput) => {
      const res = await fetch("/api/transactions", {
        method: "POST",
        headers: { "Content-Type": "application/json" },
        body: JSON.stringify(input),
      });
      if (!res.ok) throw new Error("Failed to create transaction");
      return res.json();
    },
    onMutate: async (input) => {
      // Optimistic update
      const tempId = `temp-${Date.now()}`;
      addPending(tempId);

      await queryClient.cancelQueries({ queryKey: ["transactions", input.childId] });

      const previous = queryClient.getQueryData(["transactions", input.childId]);

      queryClient.setQueryData(["transactions", input.childId], (old: any) => ({
        ...old,
        transactions: [
          {
            id: tempId,
            ...input,
            timestamp: new Date().toISOString(),
            pending: true,
          },
          ...(old?.transactions ?? []),
        ],
      }));

      return { previous, tempId };
    },
    onError: (err, input, context) => {
      // Roll back on error
      queryClient.setQueryData(["transactions", input.childId], context?.previous);
    },
    onSuccess: (data, input, context) => {
      removePending(context!.tempId);
      // Replace temp with real
      queryClient.setQueryData(["transactions", input.childId], (old: any) => ({
        ...old,
        transactions: old.transactions.map((t: any) =>
          t.id === context!.tempId ? data.transaction : t,
        ),
      }));
    },
    onSettled: (data, error, input) => {
      queryClient.invalidateQueries({ queryKey: ["transactions", input.childId] });
      queryClient.invalidateQueries({ queryKey: ["children"] }); // Balance changed
    },
  });
}""")

page_break()

# ============================================================================
# SECTION 8: SERVICES LAYER
# ============================================================================
add_heading("8. Services Layer", level=1)

add_para("Services own business rules. Each service is a stateless module that receives input, calls repositories, and returns results. The TokenService is an INTERFACE ONLY — never modify token calculations.", bold=True, size=12)

add_code("""// src/server/services/token.service.ts
// ============================================================================
// TOKEN SERVICE — BLACK-BOX INTERFACE
// ============================================================================
// DO NOT modify token pricing, incentive calculations, conversion formulas,
// or reward algorithms. These are proprietary domain rules.
// This interface is the ONLY public API for token operations.
// ============================================================================

export interface TokenService {
  /**
   * Parent awards tokens to a child.
   * The parent's account is charged at TOKEN_BUY_RATE per token.
   * Returns the transaction record.
   */
  award(input: {
    familyId: string;
    childId: string;
    tokens: number;
    note: string;
    awardedBy: string;
  }): Promise<TokenAwardResult>;

  /**
   * Child redeems tokens for cash (credited to savings).
   * Uses TOKEN_REDEEM_RATE per token.
   * Returns the transaction record + cash value.
   */
  redeem(input: {
    familyId: string;
    childId: string;
    tokens: number;
  }): Promise<TokenRedeemResult>;

  /**
   * Returns the current balance for a child.
   * Balance = awarded tokens - redeemed tokens.
   */
  getBalance(childId: string): Promise<number>;

  /**
   * Returns the cost to the parent for purchasing N tokens.
   * Uses TOKEN_BUY_RATE. This is a pure calculation — no side effects.
   */
  calculatePurchaseCost(tokens: number): number;

  /**
   * Returns the cash value a child receives for redeeming N tokens.
   * Uses TOKEN_REDEEM_RATE. This is a pure calculation — no side effects.
   */
  calculateRedeemValue(tokens: number): number;
}

export interface TokenAwardResult {
  transactionId: string;
  ledgerEntryId: string;
  tokensAwarded: number;
  parentCost: number;  // UGX
}

export interface TokenRedeemResult {
  transactionId: string;
  ledgerEntryId: string;
  tokensRedeemed: number;
  cashValue: number;  // UGX credited to savings
  newBalance: number;
}

// ============================================================================
// IMPLEMENTATION — lives in a separate file, injected via DI
// src/server/services/token.service.impl.ts
// (Contents are proprietary — do not read or modify without authorization)
// ============================================================================""")

add_code("""// src/server/services/goal.service.ts
import { GoalRepository } from "@/server/repositories/goal.repository";
import { AuditService } from "./audit.service";
import { NotificationService } from "./notification.service";
import { logger } from "@/lib/logger";

export const GoalService = {
  async create(input: CreateGoalInput, userId: string) {
    const goal = await GoalRepository.create(input);

    await AuditService.log({
      userId,
      action: "CREATE",
      entityType: "Goal",
      entityId: goal.id,
      after: goal,
    });

    logger.info({ goalId: goal.id, userId }, "Goal created");
    return goal;
  },

  async contribute(goalId: string, amount: number, contributedBy: string) {
    const goal = await GoalRepository.findById(goalId);
    if (!goal) throw new NotFoundError("Goal not found");

    // Optimistic concurrency
    const updated = await GoalRepository.updateWithVersion(
      goal.id,
      goal.version,
      {
        currentAmount: goal.currentAmount + amount,
      },
    );

    await GoalRepository.addContribution({
      goalId,
      amount,
      contributedBy,
    });

    await AuditService.log({
      userId: contributedBy,
      action: "GOAL_CONTRIBUTE",
      entityType: "Goal",
      entityId: goalId,
      after: { currentAmount: updated.currentAmount },
    });

    // Check if goal reached
    if (updated.currentAmount >= updated.targetAmount && goal.currentAmount < goal.targetAmount) {
      await NotificationService.notifyGoalReached(updated);
    }

    return updated;
  },

  async delete(goalId: string, userId: string) {
    const goal = await GoalRepository.findById(goalId);
    if (!goal) throw new NotFoundError("Goal not found");

    await AuditService.log({
      userId,
      action: "DELETE",
      entityType: "Goal",
      entityId: goalId,
      before: goal,
    });

    await GoalRepository.softDelete(goalId);
  },
};""")

page_break()

# ============================================================================
# SECTION 9: REPOSITORY LAYER
# ============================================================================
add_heading("9. Repository Layer", level=1)

add_para("Repositories own data access. They return domain entities, not Prisma rows. This isolates the ORM — if we switch from Prisma to Drizzle, only repositories change.", bold=True, size=12)

add_code("""// src/server/repositories/base.repository.ts
import { db } from "@/lib/db";

export interface BaseRepository<T, CreateInput, UpdateInput> {
  findById(id: string): Promise<T | null>;
  findMany(filter: Record<string, any>, opts?: { limit?: number; offset?: number }): Promise<T[]>;
  create(input: CreateInput): Promise<T>;
  update(id: string, input: UpdateInput): Promise<T>;
  updateWithVersion(id: string, version: number, input: Partial<UpdateInput>): Promise<T>;
  softDelete(id: string): Promise<void>;
}

// src/server/repositories/goal.repository.ts
import { db } from "@/lib/db";
import type { Goal } from "@/types/domain";

export const GoalRepository = {
  async findById(id: string): Promise<Goal | null> {
    const row = await db.goal.findUnique({
      where: { id, deletedAt: null },
      include: { contributions: true },
    });
    return row ? toDomain(row) : null;
  },

  async findMany(filter: {
    familyId?: string;
    ownerId?: string;
    visibility?: string;
  }, opts?: { limit?: number; offset?: number }): Promise<Goal[]> {
    const rows = await db.goal.findMany({
      where: {
        ...filter,
        deletedAt: null,
      },
      orderBy: { createdAt: "desc" },
      take: opts?.limit ?? 50,
      skip: opts?.offset ?? 0,
    });
    return rows.map(toDomain);
  },

  async create(input: CreateGoalInput): Promise<Goal> {
    const row = await db.goal.create({
      data: {
        ...input,
        periodStart: new Date(),
        version: 1,
      },
    });
    return toDomain(row);
  },

  async updateWithVersion(id: string, version: number, patch: Partial<Goal>): Promise<Goal> {
    try {
      const row = await db.goal.update({
        where: { id, version },  // Optimistic concurrency
        data: {
          ...patch,
          version: { increment: 1 },
        },
      });
      return toDomain(row);
    } catch (err) {
      if (err.code === "P2025") {
        throw new ConcurrencyError("Goal was modified by another user. Refresh and try again.");
      }
      throw err;
    }
  },

  async addContribution(input: {
    goalId: string;
    amount: number;
    contributedBy: string;
  }): Promise<void> {
    await db.goalContribution.create({ data: input });
  },

  async softDelete(id: string): Promise<void> {
    await db.goal.update({
      where: { id },
      data: { deletedAt: new Date() },
    });
  },
};

function toDomain(row: any): Goal {
  return {
    id: row.id,
    familyId: row.familyId,
    ownerId: row.ownerId,
    ownerKind: row.ownerKind,
    ownerName: row.ownerName,
    title: row.title,
    type: row.type,
    cadence: row.cadence,
    visibility: row.visibility,
    targetAmount: row.targetAmount,
    currentAmount: row.currentAmount,
    periodStart: row.periodStart.getTime(),
    note: row.note ?? undefined,
    createdAt: row.createdAt.getTime(),
    version: row.version,
  };
}""")

# Save
doc.save(str(OUT))
print(f"Part 2 appended. Size: {OUT.stat().st_size:,} bytes")
